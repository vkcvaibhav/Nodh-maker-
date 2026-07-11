import streamlit as st
import google.generativeai as genai
import sqlite3
import datetime
from PIL import Image
import io
import PyPDF2
from docx import Document as DocxReader
import urllib.request
import requests
import pandas as pd
import re
import os
import json

# Word Document Generation Imports
from docx import Document
from docx.shared import Mm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENTATION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Define global logo paths
NAU_LOGO = "logos/nau_logo.png"
ICAR_LOGO = "logos/icar_logo.png"

from github import Github

# ==========================================
# GitHub Cloud Sync Engine
# ==========================================
def get_secret_value(*names):
    for name in names:
        try:
            value = st.secrets.get(name)
        except Exception:
            value = None
        if value:
            return str(value).strip()
    return ""

def queue_warning(message):
    print(message)
    try:
        if st.session_state.get("_ui_ready"):
            st.warning(message)
        else:
            warnings = st.session_state.setdefault("_startup_warnings", [])
            if message not in warnings:
                warnings.append(message)
    except Exception:
        pass

def get_github_repo():
    token = get_secret_value("GITHUB_TOKEN")
    repo_name = get_secret_value("REPO_NAME") or "vkcvaibhav/Nodh-maker-"
    if not token:
        return None
    try:
        g = Github(token)
        return g.get_repo(repo_name)
    except Exception as e:
        queue_warning(f"GitHub auth failed: {e}")
        return None

def pull_db_from_github():
    """Downloads the latest DB from GitHub when the app wakes up."""
    repo = get_github_repo()
    if not repo: return
    try:
        # Assuming we store it in a 'data' folder on GitHub
        file_content = repo.get_contents(f"data/{DB_FILE}")
        with open(DB_FILE, "wb") as f:
            f.write(file_content.decoded_content)
        print("Database successfully pulled from GitHub!")
    except Exception as e:
        queue_warning(f"GitHub DB pull skipped or failed: {e}")

def push_db_to_github():
    """Uploads the local DB to GitHub after any changes."""
    repo = get_github_repo()
    if not repo:
        return False
    try:
        with open(DB_FILE, "rb") as f:
            content = f.read()
        try:
            # Update existing file
            contents = repo.get_contents(f"data/{DB_FILE}")
            repo.update_file(contents.path, f"Auto-backup DB {datetime.datetime.now()}", content, contents.sha)
        except Exception:
            # Create new file if it doesn't exist
            repo.create_file(f"data/{DB_FILE}", "Initial DB backup", content)
        return True
    except Exception as e:
        queue_warning(f"Failed to push DB to GitHub: {e}")
        return False

def push_file_to_github(file_bytes, github_path):
    """Uploads documents/PDFs to the GitHub repo."""
    repo = get_github_repo()
    if not repo:
        return False
    content = bytes(file_bytes)
    try:
        try:
            existing = repo.get_contents(github_path)
            repo.update_file(existing.path, f"Updated {github_path}", content, existing.sha)
        except Exception:
            repo.create_file(github_path, f"Uploaded {github_path}", content)
        return True
    except Exception as e:
        queue_warning(f"Failed to push file to GitHub: {e}")
        return False

def pull_file_from_github(github_path):
    """Downloads a missing document back from GitHub repository."""
    repo = get_github_repo()
    if not repo: return None
    try:
        file_content = repo.get_contents(github_path)
        return file_content.decoded_content
    except Exception as e:
        queue_warning(f"Failed to pull file from GitHub: {e}")
        return None

def load_vault_file_bytes(file_path):
    """Checks local storage first; if missing, pulls from GitHub and re-caches it."""
    if os.path.exists(file_path):
        with open(file_path, "rb") as f:
            return f.read()
    else:
        # File is missing locally due to cloud reset; pull from cloud storage
        github_path = file_path.replace("\\", "/")
        file_bytes = pull_file_from_github(github_path)
        if file_bytes:
            try:
                # Cache it locally again so subsequent calls are fast
                os.makedirs(os.path.dirname(file_path), exist_ok=True)
                with open(file_path, "wb") as f:
                    f.write(file_bytes)
            except Exception as e:
                queue_warning(f"Could not cache pulled vault file locally: {e}")
            return file_bytes
    return None
# ==========================================
# Database Setup for Archiving, Workflow & Digital Vault
# ==========================================
DB_FILE = "sadar_nondh_archive.db"

GUJARATI_DIGIT_TRANS = str.maketrans("૦૧૨૩૪૫૬૭૮૯", "0123456789")

def coerce_amount(value, default=0.0):
    """Convert DB/AI/UI currency values to a safe float."""
    try:
        default_value = float(default)
    except (TypeError, ValueError):
        default_value = 0.0

    if value is None:
        return default_value

    if isinstance(value, (int, float)) and not isinstance(value, bool):
        try:
            if pd.isna(value):
                return default_value
        except Exception:
            pass
        return float(value)

    text = str(value).strip().translate(GUJARATI_DIGIT_TRANS)
    if not text:
        return default_value

    text = text.replace(",", "")
    match = re.search(r"-?\d+(?:\.\d+)?", text)
    if not match:
        return default_value

    try:
        return float(match.group(0))
    except ValueError:
        return default_value

def format_amount(value, default=0.0):
    return f"{coerce_amount(value, default):.2f}"

def init_db():
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    # Table for Sadar Nondh
    c.execute('''CREATE TABLE IF NOT EXISTS archive 
                 (id INTEGER PRIMARY KEY AUTOINCREMENT, 
                  date TEXT, month TEXT, year TEXT, subject TEXT, content TEXT)''')
    
    # Table for Purchase Orders (Workflow Tracking) - Added nondh_id and payment_info
    c.execute('''CREATE TABLE IF NOT EXISTS purchase_orders 
                 (id INTEGER PRIMARY KEY AUTOINCREMENT, 
                  nondh_id INTEGER, vendor_name TEXT, out_no TEXT, date TEXT, amount REAL, status TEXT, payment_info TEXT)''')
                  
    # Table for Digital Vault (Uploaded PDFs/Images & Drafts) - Added nondh_id
    c.execute('''CREATE TABLE IF NOT EXISTS digital_vault
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  nondh_id INTEGER, file_name TEXT, file_path TEXT, upload_date TEXT,
                  financial_year TEXT, month TEXT, doc_type TEXT, description TEXT)''')

    c.execute('''CREATE TABLE IF NOT EXISTS app_memories
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  category TEXT, title TEXT, content TEXT, keywords TEXT,
                  priority INTEGER DEFAULT 5, active INTEGER DEFAULT 1,
                  source_type TEXT, source_id TEXT, created_at TEXT, updated_at TEXT)''')

    c.execute('''CREATE TABLE IF NOT EXISTS memory_suggestions
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  category TEXT, title TEXT, suggested_content TEXT, keywords TEXT,
                  priority INTEGER DEFAULT 5, reason TEXT, source_type TEXT, source_id TEXT,
                  source_snapshot TEXT, status TEXT DEFAULT 'Pending',
                  created_at TEXT, updated_at TEXT)''')

    c.execute('''CREATE TABLE IF NOT EXISTS app_skills
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  name TEXT, trigger_keywords TEXT, goal TEXT, steps TEXT, examples TEXT,
                  validation_rules TEXT, priority INTEGER DEFAULT 5, active INTEGER DEFAULT 1,
                  version INTEGER DEFAULT 1, created_at TEXT, updated_at TEXT)''')

    c.execute('''CREATE TABLE IF NOT EXISTS skill_suggestions
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  name TEXT, trigger_keywords TEXT, goal TEXT, steps TEXT, examples TEXT,
                  validation_rules TEXT, priority INTEGER DEFAULT 5, reason TEXT,
                  source_type TEXT, source_id TEXT, source_snapshot TEXT,
                  status TEXT DEFAULT 'Pending', created_at TEXT, updated_at TEXT)''')

    c.execute('''CREATE TABLE IF NOT EXISTS skill_runs
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  skill_id INTEGER, skill_name TEXT, workflow TEXT, context_summary TEXT,
                  outcome TEXT, source_id TEXT, created_at TEXT)''')
                  
    # Safe Database Schema Upgrades for existing users
    try: c.execute("ALTER TABLE purchase_orders ADD COLUMN nondh_id INTEGER")
    except: pass
    try: c.execute("ALTER TABLE purchase_orders ADD COLUMN payment_info TEXT")
    except: pass
    try: c.execute("ALTER TABLE digital_vault ADD COLUMN nondh_id INTEGER")
    except: pass

    conn.commit()
    conn.close()

def save_to_db(subject, content):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    now = datetime.datetime.now()
    c.execute("INSERT INTO archive (date, month, year, subject, content) VALUES (?, ?, ?, ?, ?)", 
              (now.strftime("%d/%m/%Y"), now.strftime("%m"), now.strftime("%Y"), subject, content))
    nondh_id = c.lastrowid
    conn.commit()
    conn.close()
    push_db_to_github()
    return nondh_id

def get_po_for_nondh(nondh_id):
    if not nondh_id:
        return None
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("SELECT id, status FROM purchase_orders WHERE nondh_id = ? ORDER BY id DESC LIMIT 1", (nondh_id,))
    po = c.fetchone()
    conn.close()
    return po

def save_po_to_db(nondh_id, vendor_name, out_no, date, amount):
    amount = coerce_amount(amount)
    existing_po = get_po_for_nondh(nondh_id)
    if existing_po:
        return existing_po[0]
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("INSERT INTO purchase_orders (nondh_id, vendor_name, out_no, date, amount, status) VALUES (?, ?, ?, ?, ?, 'Unfinished')", 
              (nondh_id, vendor_name, out_no, date, amount))
    po_id = c.lastrowid
    conn.commit()
    conn.close()
    push_db_to_github()
    return po_id

def get_unfinished_pos(statuses=('Unfinished',)):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    # Create the correct number of placeholders (?) for our statuses
    placeholders = ','.join(['?'] * len(statuses))
    query = f"SELECT id, nondh_id, vendor_name, out_no, date, amount FROM purchase_orders WHERE status IN ({placeholders})"
    c.execute(query, statuses)
    data = c.fetchall()
    conn.close()
    return data

def mark_po_as_payment_generated(po_id):
    """Updates the status so it disappears from Tab 4 but stays in Tab 5/6"""
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("UPDATE purchase_orders SET status = 'Payment_Generated' WHERE id = ?", (po_id,))
    conn.commit()
    push_db_to_github()
    conn.close()

def mark_po_as_paid(po_id, payment_info=""):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("UPDATE purchase_orders SET status = 'Paid', payment_info = ? WHERE id = ?", (payment_info, po_id))
    conn.commit()
    push_db_to_github()
    conn.close()
    
def delete_po(po_id):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("DELETE FROM purchase_orders WHERE id = ?", (po_id,))
    conn.commit()
    push_db_to_github()
    conn.close()

def get_archives(month, year, keyword=""):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    query = "SELECT id, date, subject, content FROM archive WHERE 1=1"
    params = []
    if year != "All":
        query += " AND year=?"
        params.append(year)
    if month != "All":
        query += " AND month=?"
        params.append(month)
    if keyword:
        query += " AND (subject LIKE ? OR content LIKE ?)"
        params.extend([f"%{keyword}%", f"%{keyword}%"])
    query += " ORDER BY id DESC"
    c.execute(query, tuple(params))
    data = c.fetchall()
    conn.close()
    return data

# --- Digital Vault Helpers ---
def get_financial_year(date_obj):
    if date_obj.month < 4: return f"{date_obj.year - 1}-{str(date_obj.year)[2:]}"
    else: return f"{date_obj.year}-{str(date_obj.year + 1)[2:]}"

def save_file_to_vault(file_bytes, original_name, doc_type, nondh_id=None, description="", upload_date=None):
    if upload_date is None: upload_date = datetime.date.today()
    fy = get_financial_year(upload_date)
    month_str = upload_date.strftime("%B")
    
    safe_fy = fy.replace("-", "_")
    folder_path = os.path.join("digital_vault", safe_fy, doc_type.replace(" ", "_"))
    os.makedirs(folder_path, exist_ok=True)
    
    timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
    safe_name = f"{timestamp}_{original_name}"
    file_path = os.path.join(folder_path, safe_name)
    
    # 1. Save locally
    with open(file_path, "wb") as f: f.write(file_bytes)
        
    # 2. Save metadata to database
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("INSERT INTO digital_vault (nondh_id, file_name, file_path, upload_date, financial_year, month, doc_type, description) VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
              (nondh_id, original_name, file_path, upload_date.strftime("%Y-%m-%d"), fy, month_str, doc_type, description))
    conn.commit()
    conn.close()
    
    # 3. Sync SQLite DB changes to GitHub
    push_db_to_github()
    
    # FIX FIX FIX: Convert windows paths if any, and push the actual file to GitHub!
    github_path = file_path.replace("\\", "/")
    push_file_to_github(file_bytes, github_path)
    
def get_vault_files_by_nondh(nondh_id):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("SELECT file_name, file_path, upload_date, doc_type, description FROM digital_vault WHERE nondh_id = ? ORDER BY id ASC", (nondh_id,))
    data = c.fetchall()
    conn.close()
    return data

def get_vault_files(fy="All", doc_type="All", search_keyword=""):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    # સુધારો: ડેટાબેઝમાંથી ડિલીટ કરવા માટે id (vault_id) ઉમેર્યો
    query = "SELECT id, nondh_id, file_name, file_path, upload_date, financial_year, month, doc_type, description FROM digital_vault WHERE 1=1"
    params = []
    
    if fy != "All":
        query += " AND financial_year=?"
        params.append(fy)
    if doc_type != "All":
        query += " AND doc_type=?"
        params.append(doc_type)
    if search_keyword:
        query += " AND (file_name LIKE ? OR description LIKE ?)"
        params.extend([f"%{search_keyword}%", f"%{search_keyword}%"])
        
    query += " ORDER BY upload_date DESC"
    c.execute(query, tuple(params))
    data = c.fetchall()
    conn.close()
    return data

# --- Learning Memory & Hermes-style Skill Helpers ---
LEARNING_CATEGORIES = [
    "nondh_style", "statute_precedent", "item_mapping", "vendor_default",
    "budget_default", "invoice_extraction", "bill_payment", "bill_pasting",
    "register_rule", "general_workflow"
]

def now_iso():
    return datetime.datetime.now().isoformat(timespec="seconds")

def compact_text(text, limit=1500):
    text = re.sub(r"\s+", " ", str(text or "")).strip()
    return text[:limit]

def redact_sensitive(text):
    text = str(text or "")
    text = re.sub(r"AIza[0-9A-Za-z_-]{20,}", "[REDACTED_GEMINI_KEY]", text)
    text = re.sub(r"gh[pousr]_[0-9A-Za-z_]{20,}", "[REDACTED_GITHUB_TOKEN]", text)
    text = re.sub(r"(?i)(api[_\s-]*key|token|secret)\s*[:=]\s*[^\s,;]+", r"\1=[REDACTED]", text)
    return text

def keyword_tokens(text):
    return set(re.findall(r"[\w\u0A80-\u0AFF]+", str(text or "").lower()))

def parse_keywords(text):
    return [kw.strip().lower() for kw in re.split(r"[,;\n]+", str(text or "")) if kw.strip()]

def learning_score(context, keywords, title="", category="", wanted_categories=None, priority=5):
    context_l = str(context or "").lower()
    tokens = keyword_tokens(context_l)
    score = int(priority or 0)
    if wanted_categories and category in wanted_categories:
        score += 6
    if category == "general_workflow":
        score += 2
    for kw in parse_keywords(keywords):
        if kw in context_l:
            score += 7
        elif kw in tokens:
            score += 4
    title_l = str(title or "").lower()
    if title_l and title_l in context_l:
        score += 4
    return score

def list_memories(include_inactive=True, status_active_only=False):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    query = "SELECT id, category, title, content, keywords, priority, active, source_type, source_id, created_at, updated_at FROM app_memories"
    if status_active_only or not include_inactive:
        query += " WHERE active = 1"
    query += " ORDER BY active DESC, priority DESC, updated_at DESC, id DESC"
    c.execute(query)
    rows = c.fetchall()
    conn.close()
    return rows

def list_memory_suggestions(status="Pending"):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("""SELECT id, category, title, suggested_content, keywords, priority, reason,
                 source_type, source_id, source_snapshot, status, created_at, updated_at
                 FROM memory_suggestions WHERE status = ? ORDER BY priority DESC, id DESC""", (status,))
    rows = c.fetchall()
    conn.close()
    return rows

def save_memory(category, title, content, keywords="", priority=5, active=1, source_type="manual", source_id=""):
    now = now_iso()
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("""INSERT INTO app_memories
                 (category, title, content, keywords, priority, active, source_type, source_id, created_at, updated_at)
                 VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
              (category, title, redact_sensitive(content), redact_sensitive(keywords), int(priority), int(active), source_type, str(source_id or ""), now, now))
    memory_id = c.lastrowid
    conn.commit()
    conn.close()
    push_db_to_github()
    return memory_id

def update_memory(memory_id, category, title, content, keywords, priority, active):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("""UPDATE app_memories SET category=?, title=?, content=?, keywords=?, priority=?,
                 active=?, updated_at=? WHERE id=?""",
              (category, title, redact_sensitive(content), redact_sensitive(keywords), int(priority), int(active), now_iso(), memory_id))
    conn.commit()
    conn.close()
    push_db_to_github()

def set_memory_active(memory_id, active):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("UPDATE app_memories SET active=?, updated_at=? WHERE id=?", (int(active), now_iso(), memory_id))
    conn.commit()
    conn.close()
    push_db_to_github()

def delete_memory(memory_id):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("DELETE FROM app_memories WHERE id=?", (memory_id,))
    conn.commit()
    conn.close()
    push_db_to_github()

def save_memory_suggestion(category, title, content, keywords="", priority=5, reason="", source_type="", source_id="", source_snapshot=""):
    now = now_iso()
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("""INSERT INTO memory_suggestions
                 (category, title, suggested_content, keywords, priority, reason, source_type,
                  source_id, source_snapshot, status, created_at, updated_at)
                 VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, 'Pending', ?, ?)""",
              (category, title, redact_sensitive(content), redact_sensitive(keywords), int(priority), reason,
               source_type, str(source_id or ""), compact_text(redact_sensitive(source_snapshot), 2500), now, now))
    suggestion_id = c.lastrowid
    conn.commit()
    conn.close()
    push_db_to_github()
    return suggestion_id

def approve_memory_suggestion(suggestion_id):
    rows = [row for row in list_memory_suggestions("Pending") if row[0] == suggestion_id]
    if not rows:
        return
    _, category, title, content, keywords, priority, _, source_type, source_id, _, _, _, _ = rows[0]
    save_memory(category, title, content, keywords, priority, 1, source_type or "suggestion", source_id)
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("UPDATE memory_suggestions SET status='Approved', updated_at=? WHERE id=?", (now_iso(), suggestion_id))
    conn.commit()
    conn.close()
    push_db_to_github()

def reject_memory_suggestion(suggestion_id):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("UPDATE memory_suggestions SET status='Rejected', updated_at=? WHERE id=?", (now_iso(), suggestion_id))
    conn.commit()
    conn.close()
    push_db_to_github()

def list_skills(include_inactive=True):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    query = """SELECT id, name, trigger_keywords, goal, steps, examples, validation_rules,
               priority, active, version, created_at, updated_at FROM app_skills"""
    if not include_inactive:
        query += " WHERE active = 1"
    query += " ORDER BY active DESC, priority DESC, updated_at DESC, id DESC"
    c.execute(query)
    rows = c.fetchall()
    conn.close()
    return rows

def save_skill(name, trigger_keywords, goal, steps, examples="", validation_rules="", priority=5, active=1, version=1):
    now = now_iso()
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("""INSERT INTO app_skills
                 (name, trigger_keywords, goal, steps, examples, validation_rules, priority, active, version, created_at, updated_at)
                 VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
              (name, redact_sensitive(trigger_keywords), redact_sensitive(goal), redact_sensitive(steps),
               redact_sensitive(examples), redact_sensitive(validation_rules), int(priority), int(active), int(version), now, now))
    skill_id = c.lastrowid
    conn.commit()
    conn.close()
    push_db_to_github()
    return skill_id

def update_skill(skill_id, name, trigger_keywords, goal, steps, examples, validation_rules, priority, active, version):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("""UPDATE app_skills SET name=?, trigger_keywords=?, goal=?, steps=?, examples=?,
                 validation_rules=?, priority=?, active=?, version=?, updated_at=? WHERE id=?""",
              (name, redact_sensitive(trigger_keywords), redact_sensitive(goal), redact_sensitive(steps),
               redact_sensitive(examples), redact_sensitive(validation_rules), int(priority), int(active), int(version), now_iso(), skill_id))
    conn.commit()
    conn.close()
    push_db_to_github()

def set_skill_active(skill_id, active):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("UPDATE app_skills SET active=?, updated_at=? WHERE id=?", (int(active), now_iso(), skill_id))
    conn.commit()
    conn.close()
    push_db_to_github()

def delete_skill(skill_id):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("DELETE FROM app_skills WHERE id=?", (skill_id,))
    conn.commit()
    conn.close()
    push_db_to_github()

def list_skill_suggestions(status="Pending"):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("""SELECT id, name, trigger_keywords, goal, steps, examples, validation_rules,
                 priority, reason, source_type, source_id, source_snapshot, status, created_at, updated_at
                 FROM skill_suggestions WHERE status = ? ORDER BY priority DESC, id DESC""", (status,))
    rows = c.fetchall()
    conn.close()
    return rows

def save_skill_suggestion(name, trigger_keywords, goal, steps, examples="", validation_rules="", priority=5, reason="", source_type="", source_id="", source_snapshot=""):
    now = now_iso()
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("""INSERT INTO skill_suggestions
                 (name, trigger_keywords, goal, steps, examples, validation_rules, priority, reason,
                  source_type, source_id, source_snapshot, status, created_at, updated_at)
                 VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'Pending', ?, ?)""",
              (name, redact_sensitive(trigger_keywords), redact_sensitive(goal), redact_sensitive(steps),
               redact_sensitive(examples), redact_sensitive(validation_rules), int(priority), reason,
               source_type, str(source_id or ""), compact_text(redact_sensitive(source_snapshot), 2500), now, now))
    suggestion_id = c.lastrowid
    conn.commit()
    conn.close()
    push_db_to_github()
    return suggestion_id

def approve_skill_suggestion(suggestion_id):
    rows = [row for row in list_skill_suggestions("Pending") if row[0] == suggestion_id]
    if not rows:
        return
    _, name, trigger_keywords, goal, steps, examples, validation_rules, priority, *_ = rows[0]
    save_skill(name, trigger_keywords, goal, steps, examples, validation_rules, priority, 1, 1)
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("UPDATE skill_suggestions SET status='Approved', updated_at=? WHERE id=?", (now_iso(), suggestion_id))
    conn.commit()
    conn.close()
    push_db_to_github()

def reject_skill_suggestion(suggestion_id):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("UPDATE skill_suggestions SET status='Rejected', updated_at=? WHERE id=?", (now_iso(), suggestion_id))
    conn.commit()
    conn.close()
    push_db_to_github()

def get_relevant_memories(context, categories=None, limit=8):
    wanted = set(categories or [])
    candidates = []
    for row in list_memories(include_inactive=False):
        memory_id, category, title, content, keywords, priority, active, *_ = row
        if wanted and category not in wanted and category != "general_workflow":
            continue
        score = learning_score(context, keywords, title, category, wanted, priority)
        if score > 0:
            candidates.append({"id": memory_id, "category": category, "title": title, "content": content, "keywords": keywords, "priority": priority, "score": score})
    return sorted(candidates, key=lambda x: (x["score"], x["priority"], x["id"]), reverse=True)[:limit]

def get_relevant_skills(context, workflow="", limit=3):
    candidates = []
    combined_context = f"{workflow} {context}"
    for row in list_skills(include_inactive=False):
        skill_id, name, trigger_keywords, goal, steps, examples, validation_rules, priority, active, version, *_ = row
        score = learning_score(combined_context, trigger_keywords, name, "skill", None, priority)
        if score > 0:
            candidates.append({
                "id": skill_id, "name": name, "trigger_keywords": trigger_keywords, "goal": goal,
                "steps": steps, "examples": examples, "validation_rules": validation_rules,
                "priority": priority, "version": version, "score": score
            })
    return sorted(candidates, key=lambda x: (x["score"], x["priority"], x["id"]), reverse=True)[:limit]

def build_learning_prompt(memories, skills):
    sections = []
    if memories:
        lines = ["Approved Memories:"]
        for m in memories:
            lines.append(f"- [{m['category']} | priority {m['priority']}] {m['title']}: {compact_text(m['content'], 700)}")
        sections.append("\n".join(lines))
    if skills:
        lines = ["Approved Hermes-Style Process Skills:"]
        for s in skills:
            lines.append(
                f"- Skill: {s['name']} v{s['version']} (priority {s['priority']})\n"
                f"  Goal: {compact_text(s['goal'], 350)}\n"
                f"  Steps: {compact_text(s['steps'], 700)}\n"
                f"  Validation: {compact_text(s['validation_rules'], 350)}"
            )
        sections.append("\n".join(lines))
    return "\n\n".join(sections)

def get_learning_context(context, categories=None, workflow="", memory_limit=8, skill_limit=3):
    try:
        safe_context = redact_sensitive(compact_text(context, 3000))
        memories = get_relevant_memories(safe_context, categories, memory_limit)
        skills = get_relevant_skills(safe_context, workflow, skill_limit)
        return {"memories": memories, "skills": skills, "prompt": build_learning_prompt(memories, skills)}
    except Exception as e:
        queue_warning(f"Learning memory retrieval failed: {e}")
        return {"memories": [], "skills": [], "prompt": ""}

def log_skill_runs(skills, workflow, context_summary="", outcome="used", source_id=""):
    if not skills:
        return
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    for skill in skills:
        c.execute("""INSERT INTO skill_runs (skill_id, skill_name, workflow, context_summary, outcome, source_id, created_at)
                     VALUES (?, ?, ?, ?, ?, ?, ?)""",
                  (skill.get("id"), skill.get("name"), workflow, compact_text(redact_sensitive(context_summary), 900), outcome, str(source_id or ""), now_iso()))
    conn.commit()
    conn.close()
    push_db_to_github()

def list_skill_runs(limit=50):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("""SELECT id, skill_name, workflow, context_summary, outcome, source_id, created_at
                 FROM skill_runs ORDER BY id DESC LIMIT ?""", (limit,))
    rows = c.fetchall()
    conn.close()
    return rows

def seed_default_skills():
    defaults = [
        ("Chemical Purchase Nondh", "chemical, chemicals, reagent, laboratory, acid, solvent, acarology",
         "Create a purchase Nondh for chemicals or laboratory consumables.",
         "Identify item names, quantities, pack sizes, and unit prices. Use English table headers. Keep package size only in Available Pkt/Unit. Use Statute 121 precedent for research/lab materials.",
         "Acetic Acid, Ethanol, laboratory reagents.",
         "Total Price must equal Required Quantity x Unit/Pkt Price. Do not invent package units.", 8),
        ("Equipment Purchase Nondh", "equipment, instrument, device, microscope, machine, repair, purchase",
         "Create a purchase Nondh for equipment or instruments.",
         "Clarify whether the item is new equipment or repair/maintenance. Choose statute precedent based on equipment category. Keep justification practical and formal.",
         "Microscope accessories, lab instruments.",
         "Do not classify consumables as equipment. Mention scheme/budget head consistently.", 7),
        ("Statute 121 Item Matching", "statute, 121, item number, approval, sanction, precedent",
         "Select the most defensible Statute 121 item number using sample precedent first.",
         "Check sample Nondh precedent. If no match exists, use statute PDF context. Explain chosen item and reject a plausible wrong item.",
         "Use exact item format like ૫૪ (i) when precedent supports it.",
         "Never hallucinate item numbers. Always include justification and rejected alternative.", 10),
        ("Invoice Total Extraction", "invoice, bill, total, payable, gst, net payable, grand total",
         "Extract the final payable amount and bill number from uploaded invoices.",
         "Prefer Grand Total, Invoice Total, Net Payable, or Total Amount including taxes. Compare against PO amount and return clean JSON.",
         "Vendor bill PDF/image extraction.",
         "Return pure numeric amount without commas. Do not choose subtotal when tax-inclusive total exists.", 8),
        ("Bill Pasting Register Selection", "bill pasting, register, stock, consumable, deadstock, register page",
         "Guide bill pasting register details and certificate wording.",
         "Use consumable register for consumables, deadstock for durable equipment, and stock register where office stock entry is required.",
         "Chemicals usually map to consumable usage; durable instruments often map to deadstock.",
         "Do not auto-fill page numbers unless provided by the user.", 7),
    ]
    existing_names = {str(row[1]).lower() for row in list_skills(include_inactive=True)}
    for name, keywords, goal, steps, examples, validation, priority in defaults:
        if name.lower() not in existing_names:
            save_skill(name, keywords, goal, steps, examples, validation, priority, 1, 1)

def suggest_learning_from_nondh(api_key, original_doc, final_doc, subject, nondh_id):
    if not api_key:
        return 0
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-3.1-flash-lite-preview')
        prompt = f"""
        You are helping improve a Gujarati administrative Streamlit app.
        Compare the AI draft and final saved Nondh. Suggest reusable learning ONLY if it would help future workflows.
        Do not include API keys, secrets, or personal credentials.

        Return ONLY valid JSON:
        {{
          "memories": [
            {{"category": "nondh_style", "title": "...", "content": "...", "keywords": "comma,list", "priority": 5, "reason": "..."}}
          ],
          "skills": [
            {{"name": "...", "trigger_keywords": "comma,list", "goal": "...", "steps": "...", "examples": "...", "validation_rules": "...", "priority": 5, "reason": "..."}}
          ]
        }}

        Allowed memory categories: {", ".join(LEARNING_CATEGORIES)}
        Subject: {redact_sensitive(subject)}
        AI draft:
        {compact_text(redact_sensitive(original_doc), 5000)}

        Final saved document:
        {compact_text(redact_sensitive(final_doc), 5000)}
        """
        response = model.generate_content(prompt)
        raw = response.text.strip().replace("```json", "").replace("```", "")
        data = json.loads(raw)
        saved = 0
        snapshot = f"Subject: {subject}\nFinal: {compact_text(final_doc, 1800)}"
        for item in data.get("memories", [])[:5]:
            category = item.get("category", "general_workflow")
            if category not in LEARNING_CATEGORIES:
                category = "general_workflow"
            save_memory_suggestion(
                category, item.get("title", "Suggested memory"), item.get("content", ""),
                item.get("keywords", ""), int(item.get("priority", 5)), item.get("reason", ""),
                "nondh", nondh_id, snapshot
            )
            saved += 1
        for item in data.get("skills", [])[:3]:
            save_skill_suggestion(
                item.get("name", "Suggested skill"), item.get("trigger_keywords", ""),
                item.get("goal", ""), item.get("steps", ""), item.get("examples", ""),
                item.get("validation_rules", ""), int(item.get("priority", 5)), item.get("reason", ""),
                "nondh", nondh_id, snapshot
            )
            saved += 1
        return saved
    except Exception as e:
        queue_warning(f"Learning suggestion generation failed: {e}")
        return 0

DASHBOARD_REQUIRED_DOCS = [
    "Signed Nondh",
    "Signed PO",
    "Party Invoice",
    "Signed Bill Payment",
    "Signed Bill Pasting",
]

DASHBOARD_MONTHS = [
    "January", "February", "March", "April", "May", "June",
    "July", "August", "September", "October", "November", "December",
]

def parse_dashboard_date(value):
    """Parse common DB/UI date strings without throwing."""
    if isinstance(value, datetime.datetime):
        return value.date()
    if isinstance(value, datetime.date):
        return value
    text = str(value or "").strip()
    if not text:
        return None
    text = text.translate(GUJARATI_DIGIT_TRANS)
    for fmt in ("%d/%m/%Y", "%d.%m.%Y", "%Y-%m-%d", "%d-%m-%Y", "%d/%m/%y", "%d.%m.%y"):
        try:
            return datetime.datetime.strptime(text, fmt).date()
        except ValueError:
            pass
    parts = re.findall(r"\d+", text)
    if len(parts) < 3:
        return None
    try:
        if len(parts[0]) == 4:
            year, month, day = int(parts[0]), int(parts[1]), int(parts[2])
        else:
            day, month, year = int(parts[0]), int(parts[1]), int(parts[2])
            if year < 100:
                year += 2000
        return datetime.date(year, month, day)
    except Exception:
        return None

def date_from_archive_fields(date_value, month_value, year_value):
    parsed = parse_dashboard_date(date_value)
    if parsed:
        return parsed
    try:
        month_num = int(str(month_value or "").translate(GUJARATI_DIGIT_TRANS))
        year_num = int(str(year_value or "").translate(GUJARATI_DIGIT_TRANS))
        return datetime.date(year_num, month_num, 1)
    except Exception:
        return None

def age_label(date_obj):
    if not date_obj:
        return "unknown date"
    days = (datetime.date.today() - date_obj).days
    if days < 0:
        return "future"
    if days <= 2:
        return "fresh"
    if days <= 10:
        return "waiting"
    return "overdue"

def dashboard_date_text(date_obj):
    return date_obj.strftime("%Y-%m-%d") if date_obj else ""

def month_filter_matches(date_obj, month_text):
    if month_text == "All":
        return True
    if not date_obj:
        return False
    return date_obj.strftime("%B") == month_text

def fy_filter_matches(date_obj, fy_text):
    if fy_text == "All":
        return True
    if not date_obj:
        return False
    return get_financial_year(date_obj) == fy_text

def matches_dashboard_filters(date_obj, fy_filter="All", month_filter="All"):
    return fy_filter_matches(date_obj, fy_filter) and month_filter_matches(date_obj, month_filter)

def fetch_dashboard_rows():
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    data = {"archives": [], "purchase_orders": [], "vault": [], "errors": []}
    try:
        c.execute("SELECT id, date, month, year, subject, content FROM archive ORDER BY id DESC")
        data["archives"] = c.fetchall()
    except Exception as e:
        data["errors"].append(f"Could not read archive table: {e}")
    try:
        c.execute("""SELECT id, nondh_id, vendor_name, out_no, date, amount, status, payment_info
                     FROM purchase_orders ORDER BY id DESC""")
        data["purchase_orders"] = c.fetchall()
    except Exception as e:
        data["errors"].append(f"Could not read purchase_orders table: {e}")
    try:
        c.execute("""SELECT id, nondh_id, file_name, file_path, upload_date, financial_year, month, doc_type, description
                     FROM digital_vault ORDER BY id DESC""")
        data["vault"] = c.fetchall()
    except Exception as e:
        data["errors"].append(f"Could not read digital_vault table: {e}")
    conn.close()
    return data

def count_pending_learning_items():
    counts = {"memories": 0, "skills": 0}
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    try:
        c.execute("SELECT COUNT(*) FROM memory_suggestions WHERE status = 'Pending'")
        counts["memories"] = int(c.fetchone()[0] or 0)
    except Exception:
        counts["memories"] = 0
    try:
        c.execute("SELECT COUNT(*) FROM skill_suggestions WHERE status = 'Pending'")
        counts["skills"] = int(c.fetchone()[0] or 0)
    except Exception:
        counts["skills"] = 0
    conn.close()
    return counts

def build_dashboard_data(fy_filter="All", month_filter="All"):
    raw = fetch_dashboard_rows()
    learning_counts = count_pending_learning_items()

    archives = []
    for row in raw["archives"]:
        nondh_id, date_text, month_text, year_text, subject, content = row
        date_obj = date_from_archive_fields(date_text, month_text, year_text)
        if matches_dashboard_filters(date_obj, fy_filter, month_filter):
            archives.append({
                "id": nondh_id, "date": date_obj, "date_text": date_text,
                "subject": subject or "", "content": content or "",
            })

    all_pos = []
    filtered_pos = []
    for row in raw["purchase_orders"]:
        po_id, nondh_id, vendor, out_no, date_text, amount, status, payment_info = row
        date_obj = parse_dashboard_date(date_text)
        po = {
            "id": po_id, "nondh_id": nondh_id, "vendor": vendor or "", "out_no": out_no or "",
            "date": date_obj, "date_text": date_text or "", "amount": coerce_amount(amount),
            "status": status or "Unfinished", "payment_info": payment_info or "",
        }
        all_pos.append(po)
        if matches_dashboard_filters(date_obj, fy_filter, month_filter):
            filtered_pos.append(po)

    vault_by_nondh = {}
    all_vault_records = []
    filtered_vault_records = []
    for row in raw["vault"]:
        vault_id, nondh_id, file_name, file_path, upload_date, fy, month, doc_type, description = row
        upload_dt = parse_dashboard_date(upload_date)
        record = {
            "id": vault_id, "nondh_id": nondh_id, "file_name": file_name or "",
            "file_path": file_path or "", "upload_date": upload_dt, "financial_year": fy or "",
            "month": month or "", "doc_type": doc_type or "", "description": description or "",
        }
        all_vault_records.append(record)
        if nondh_id:
            vault_by_nondh.setdefault(nondh_id, {}).setdefault(doc_type or "", []).append(record)
        fy_ok = fy_filter == "All" or fy == fy_filter or fy_filter_matches(upload_dt, fy_filter)
        month_ok = month_filter == "All" or month == month_filter or month_filter_matches(upload_dt, month_filter)
        if fy_ok and month_ok:
            filtered_vault_records.append(record)

    archive_by_id = {item["id"]: item for item in archives}
    all_archive_ids = {row[0] for row in raw["archives"]}
    all_po_by_nondh = {}
    for po in all_pos:
        if po["nondh_id"]:
            all_po_by_nondh.setdefault(po["nondh_id"], []).append(po)

    po_by_nondh = {}
    for po in filtered_pos:
        if po["nondh_id"]:
            po_by_nondh.setdefault(po["nondh_id"], []).append(po)

    workflow_keys = []
    for archive in archives:
        workflow_keys.append(("nondh", archive["id"]))
    for po in filtered_pos:
        if po["nondh_id"]:
            key = ("nondh", po["nondh_id"])
        else:
            key = ("po", po["id"])
        if key not in workflow_keys:
            workflow_keys.append(key)

    document_rows = []
    missing_document_total = 0
    for key_type, key_id in workflow_keys:
        nondh_id = key_id if key_type == "nondh" else None
        related_pos = po_by_nondh.get(nondh_id, []) if nondh_id else [po for po in filtered_pos if po["id"] == key_id]
        latest_po = related_pos[0] if related_pos else None
        archive = archive_by_id.get(nondh_id) if nondh_id else None
        existing_docs = set(vault_by_nondh.get(nondh_id, {}).keys()) if nondh_id else set()
        missing_docs = [doc for doc in DASHBOARD_REQUIRED_DOCS if doc not in existing_docs]
        missing_document_total += len(missing_docs)
        document_rows.append({
            "Nondh ID": nondh_id or "",
            "PO ID": latest_po["id"] if latest_po else "",
            "Subject / Vendor": archive["subject"] if archive else (latest_po["vendor"] if latest_po else ""),
            "Signed Nondh": "Yes" if "Signed Nondh" in existing_docs else "Missing",
            "Signed PO": "Yes" if "Signed PO" in existing_docs else "Missing",
            "Party Invoice": "Yes" if "Party Invoice" in existing_docs else "Missing",
            "Signed Bill Payment": "Yes" if "Signed Bill Payment" in existing_docs else "Missing",
            "Signed Bill Pasting": "Yes" if "Signed Bill Pasting" in existing_docs else "Missing",
            "Missing Items": ", ".join(missing_docs),
            "Missing Count": len(missing_docs),
        })

    next_actions = []
    for archive in archives:
        if not all_po_by_nondh.get(archive["id"]):
            next_actions.append({
                "Priority": "High" if age_label(archive["date"]) == "overdue" else "Normal",
                "Workflow": "Nondh saved but no PO",
                "Nondh ID": archive["id"],
                "PO ID": "",
                "Subject / Vendor": archive["subject"],
                "Action": "Create purchase order or close if not required",
                "Age": age_label(archive["date"]),
                "Date": dashboard_date_text(archive["date"]),
                "Amount": "",
            })

    for po in filtered_pos:
        existing_docs = set(vault_by_nondh.get(po["nondh_id"], {}).keys()) if po["nondh_id"] else set()
        subject = archive_by_id.get(po["nondh_id"], {}).get("subject", po["vendor"]) if po["nondh_id"] else po["vendor"]
        if "Party Invoice" not in existing_docs and po["status"] in ("Unfinished", "Payment_Generated"):
            next_actions.append({
                "Priority": "High" if age_label(po["date"]) == "overdue" else "Normal",
                "Workflow": "PO created but party invoice missing",
                "Nondh ID": po["nondh_id"] or "",
                "PO ID": po["id"],
                "Subject / Vendor": subject,
                "Action": "Upload party invoice in Bill Payment",
                "Age": age_label(po["date"]),
                "Date": dashboard_date_text(po["date"]),
                "Amount": format_amount(po["amount"]),
            })
        if po["status"] == "Payment_Generated" and "Signed Bill Pasting" not in existing_docs:
            next_actions.append({
                "Priority": "High",
                "Workflow": "Payment form generated but bill pasting missing",
                "Nondh ID": po["nondh_id"] or "",
                "PO ID": po["id"],
                "Subject / Vendor": subject,
                "Action": "Generate/upload signed bill pasting form",
                "Age": age_label(po["date"]),
                "Date": dashboard_date_text(po["date"]),
                "Amount": format_amount(po["amount"]),
            })
        signed_missing = [doc for doc in ("Signed Nondh", "Signed PO", "Signed Bill Payment", "Signed Bill Pasting") if doc not in existing_docs]
        if signed_missing:
            next_actions.append({
                "Priority": "Normal",
                "Workflow": "Signed documents missing in vault",
                "Nondh ID": po["nondh_id"] or "",
                "PO ID": po["id"],
                "Subject / Vendor": subject,
                "Action": "Upload: " + ", ".join(signed_missing),
                "Age": age_label(po["date"]),
                "Date": dashboard_date_text(po["date"]),
                "Amount": format_amount(po["amount"]),
            })

    if learning_counts["memories"] or learning_counts["skills"]:
        next_actions.append({
            "Priority": "Normal",
            "Workflow": "Learning suggestions waiting for approval",
            "Nondh ID": "",
            "PO ID": "",
            "Subject / Vendor": "Memory and skill queue",
            "Action": f"Review {learning_counts['memories']} memory and {learning_counts['skills']} skill suggestions",
            "Age": "waiting",
            "Date": "",
            "Amount": "",
        })

    status_order = {"High": 0, "Normal": 1}
    next_actions = sorted(next_actions, key=lambda row: (status_order.get(row["Priority"], 2), row["Age"] != "overdue", row["Date"]))

    pending_payments = []
    for po in filtered_pos:
        if po["status"] != "Paid":
            pending_payments.append({
                "PO ID": po["id"],
                "Nondh ID": po["nondh_id"] or "",
                "Vendor": po["vendor"],
                "Outward No": po["out_no"],
                "Status": po["status"],
                "Date": dashboard_date_text(po["date"]),
                "Age": age_label(po["date"]),
                "Amount": format_amount(po["amount"]),
            })

    metrics = {
        "total_nondhs": len(archives),
        "nondhs_awaiting_po": sum(1 for item in archives if not all_po_by_nondh.get(item["id"])),
        "unfinished_pos": sum(1 for po in filtered_pos if po["status"] == "Unfinished"),
        "payment_forms_generated": sum(1 for po in filtered_pos if po["status"] == "Payment_Generated"),
        "paid_bills": sum(1 for po in filtered_pos if po["status"] == "Paid"),
        "pending_amount": sum(po["amount"] for po in filtered_pos if po["status"] != "Paid"),
        "paid_amount": sum(po["amount"] for po in filtered_pos if po["status"] == "Paid"),
        "missing_documents": missing_document_total,
        "pending_memory_suggestions": learning_counts["memories"],
        "pending_skill_suggestions": learning_counts["skills"],
        "vault_documents": len(filtered_vault_records),
    }

    return {
        "filters": {"financial_year": fy_filter, "month": month_filter},
        "metrics": metrics,
        "next_actions": next_actions,
        "document_rows": document_rows,
        "pending_payments": pending_payments,
        "money_summary": [
            {"Bucket": "Pending", "Count": sum(1 for po in filtered_pos if po["status"] != "Paid"), "Amount": format_amount(metrics["pending_amount"])},
            {"Bucket": "Paid", "Count": metrics["paid_bills"], "Amount": format_amount(metrics["paid_amount"])},
        ],
        "learning_queue": [
            {"Queue": "Pending memory suggestions", "Count": learning_counts["memories"]},
            {"Queue": "Pending skill suggestions", "Count": learning_counts["skills"]},
        ],
        "errors": raw["errors"],
    }

def build_dashboard_coach_prompt(dashboard_data):
    metrics = dashboard_data.get("metrics", {})
    payload = {
        "filters": dashboard_data.get("filters", {}),
        "metrics": metrics,
        "top_next_actions": dashboard_data.get("next_actions", [])[:12],
        "document_gaps": [row for row in dashboard_data.get("document_rows", []) if row.get("Missing Count", 0) > 0][:12],
        "pending_payments": dashboard_data.get("pending_payments", [])[:12],
        "learning_queue": dashboard_data.get("learning_queue", []),
    }
    return redact_sensitive(json.dumps(payload, ensure_ascii=False, indent=2))

def delete_vault_record(vault_id, nondh_id, file_path):
    """વોલ્ટમાંથી ફાઈલ કાઢે છે, અને જો Nondh જોડાયેલી હોય તો તેને પણ કાઢીને GitHub પર સિન્ક કરે છે."""
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    
    # ૧. વોલ્ટ ડેટાબેઝમાંથી ડિલીટ
    c.execute("DELETE FROM digital_vault WHERE id = ?", (vault_id,))
    
    # ૨. જો Nondh ID હોય, તો આર્કાઇવ અને પેમેન્ટ ઓર્ડરમાંથી પણ આખી Nondh કાઢી નાખો
    if nondh_id:
        c.execute("DELETE FROM archive WHERE id = ?", (nondh_id,))
        c.execute("DELETE FROM purchase_orders WHERE nondh_id = ?", (nondh_id,))
        
    conn.commit()
    conn.close()
    
    # ૩. ડેટાબેઝ GitHub પર અપડેટ કરો (જેથી ત્યાંથી પણ ડિલીટ થઈ જાય)
    push_db_to_github()
    
    # ૪. લોકલ ફોલ્ડરમાંથી PDF/Word ફાઈલ કાઢી નાખો
    if os.path.exists(file_path):
        try:
            os.remove(file_path)
        except:
            pass

# Check if we already pulled the DB in this session to avoid constant downloading
if "db_synced" not in st.session_state:
    pull_db_from_github()
    st.session_state.db_synced = True

init_db()
try:
    seed_default_skills()
except Exception as e:
    queue_warning(f"Default skill seeding failed: {e}")

def get_recent_nondhs(days=30):
    """છેલ્લા 30 દિવસની સાદર નોંધ મેળવવા માટે"""
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("SELECT id, date, subject FROM archive ORDER BY id DESC LIMIT 50")
    records = c.fetchall()
    conn.close()
    
    recent_records = []
    cutoff_date = datetime.date.today() - datetime.timedelta(days=days)
    for row in records:
        try:
            # ડેટાબેઝમાં તારીખ DD/MM/YYYY ફોર્મેટમાં છે
            row_date = datetime.datetime.strptime(row[1], "%d/%m/%Y").date()
            if row_date >= cutoff_date:
                recent_records.append(row)
        except Exception:
            pass # જો તારીખમાં ભૂલ હોય તો સ્કીપ કરો
    return recent_records

def delete_nondh(nondh_id):
    """સાદર નોંધને ડેટાબેઝમાંથી કાયમ માટે ડિલીટ કરવા માટે"""
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("DELETE FROM archive WHERE id = ?", (nondh_id,))
    conn.commit()
    conn.close()
    push_db_to_github()
init_db()

# ==========================================
# Permanent Attachments & Parsing (GitHub)
# ==========================================
@st.cache_data(ttl=3600) 
def load_permanent_context():
    statute_text = "Statute 121 Rules:\n"
    sample_text = "Sample Nondh Format:\n"
    pdf_url = "https://raw.githubusercontent.com/vkcvaibhav/Nodh-maker-/main/121_Statutes.pdf"
    docx_url = "https://raw.githubusercontent.com/vkcvaibhav/Nodh-maker-/main/sample_nondh.docx"
    try:
        r_pdf = requests.get(pdf_url)
        if r_pdf.status_code == 200:
            f = io.BytesIO(r_pdf.content)
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages: statute_text += page.extract_text() + "\n"
    except Exception: pass
    try:
        r_docx = requests.get(docx_url)
        if r_docx.status_code == 200:
            f = io.BytesIO(r_docx.content)
            doc = DocxReader(f)
            for para in doc.paragraphs: sample_text += para.text + "\n"
    except Exception: pass
    return statute_text, sample_text

def search_sample_nondh(keyword, month, year):
    try: _, sample_text = load_permanent_context()
    except Exception: return []
    if not sample_text: return []
    guj_to_eng = str.maketrans("૦૧૨૩૪૫૬૭૮૯", "0123456789")
    blocks = re.split(r'\n(?=તા\.\s*)', '\n' + sample_text)
    results = []
    for block in blocks:
        block = block.strip()
        if not block or "સાદર નોંધ" not in block: continue
        eng_block = block.translate(guj_to_eng)
        if keyword and keyword.lower() not in block.lower() and keyword.lower() not in eng_block.lower(): continue
        date_match = re.search(r'તા\.\s*([\d/ \-]+)', eng_block)
        date_str = date_match.group(1).strip() if date_match else "Unknown"
        if year != "All" and year not in date_str: continue
        if month != "All":
            if f"/{month}/" not in date_str and f"-{month}-" not in date_str and f"{month}/" not in date_str: continue
        sub_match = re.search(r'વિષય:\s*([^\n]+)', block)
        subject_str = sub_match.group(1).strip() if sub_match else "Historical Reference"
        orig_date_match = re.search(r'તા\.\s*([^\n]+)', block)
        display_date = orig_date_match.group(1).strip() if orig_date_match else "Unknown"
        results.append((None, display_date + " [Old Sample Ref]", subject_str, block))
    return results

def parse_markdown_to_parts(text):
    lines = text.split('\n')
    pre_text, table_lines, post_text = [], [], []
    in_table, table_done = False, False
    for line in lines:
        if line.strip().startswith('|'):
            in_table = True
            table_lines.append(line)
        else:
            if in_table:
                table_done = True
                in_table = False
            if not table_done: pre_text.append(line)
            else: post_text.append(line)
    df = pd.DataFrame()
    if table_lines:
        parts = table_lines[0].split('|')
        if len(parts) > 2:
            header = [x.strip() for x in parts[1:-1]]
            data = []
            for line in table_lines[2:]: 
                row_parts = line.split('|')
                if len(row_parts) > 2:
                    data.append([x.strip() for x in row_parts[1:-1]])
            if data:
                df = pd.DataFrame(data, columns=header)
                if 'Details' in df.columns: df = df[~df['Details'].astype(str).str.contains('Grand Total', case=False, na=False)]
    return "\n".join(pre_text), df, "\n".join(post_text)

def df_to_markdown_with_total(df):
    if df.empty: return ""
    grand_total = 0
    if 'Total Price' in df.columns:
        grand_total = pd.to_numeric(df['Total Price'], errors='coerce').fillna(0).sum()
    markdown = "|" + "|".join(df.columns) + "|\n|" + "|".join(["---"] * len(df.columns)) + "|\n"
    for _, row in df.iterrows():
        clean_row = [str(int(x)) if isinstance(x, float) and x.is_integer() else str(x) for x in row]
        markdown += "|" + "|".join(clean_row) + "|\n"
    if len(df.columns) >= 5:
        total_row = [""] * len(df.columns)
        total_row[1] = "**Grand Total**"
        total_row[-1] = f"**{grand_total:.2f}**"
        markdown += "|" + "|".join(total_row) + "|\n"
    return markdown

def add_bottom_border(paragraph, size='24'):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ==========================================
# Word Generators
# ==========================================
def create_docx(content):
    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Mm(210), Mm(297)
    section.left_margin, section.right_margin = Mm(42), Mm(15)
    section.top_margin, section.bottom_margin = Mm(15), Mm(15)
    
    font = doc.styles['Normal'].font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Shruti')
    font._element.append(rFonts)

    lines = content.split('\n')
    table_data = []
    in_table = False
    sig_buffer = []

    def flush_signatures():
        if sig_buffer:
            doc.add_paragraph().paragraph_format.space_before = Pt(20)
            sig_table = doc.add_table(rows=1, cols=3)
            for c in sig_table.columns: 
                for cell in c.cells: cell.width = Mm(51)
            for i, sig in enumerate(sig_buffer):
                if i < 3:
                    p = sig_table.cell(0, i).paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    parts = sig.split(',')
                    for j, part in enumerate(parts):
                        run = p.add_run(part.strip())
                        if j < len(parts) - 1: run.add_break()
            sig_buffer.clear()

    def build_and_format_table(data):
        num_cols = len(data[0])
        table = doc.add_table(rows=len(data), cols=num_cols)
        table.style = 'Table Grid'
        
        # Turn off autofit to force Word to respect your explicit column widths
        table.autofit = False 
        
        # Calculated to fit exactly within your page margins (153mm usable width)
        # Sr. No (10mm) + Details (75mm) + Qty (17mm) + Avail (17mm) + Unit Price (17mm) + Total (17mm) = 153mm
        widths = [Mm(10), Mm(75), Mm(17), Mm(17), Mm(17), Mm(17)]
        
        for row_idx, row_data in enumerate(data):
            row_cells = table.rows[row_idx].cells
            for col_idx, cell_text in enumerate(row_data):
                cell = row_cells[col_idx]
                
                # Apply width to every cell individually to enforce it in MS Word
                if num_cols == 6: 
                    cell.width = widths[col_idx]
                    
                is_bold = (row_idx == 0) or ('**' in cell_text)
                cell.text = cell_text.replace('**', '')
                p = cell.paragraphs[0]
                
                # Center align all columns EXCEPT the Details column (col_idx == 1)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if (col_idx == 1 and row_idx > 0) else WD_ALIGN_PARAGRAPH.CENTER
                
                if is_bold:
                    for run in p.runs: run.bold = True
        doc.add_paragraph()

    for line in lines:
        line_stripped = line.strip()
        if not line_stripped: continue
        if line_stripped.startswith('|'):
            in_table = True
            if not line_stripped.replace('|', '').replace('-', '').replace(' ', ''): continue
            parts = line_stripped.split('|')
            if len(parts) > 2: table_data.append([cell.strip() for cell in parts[1:-1]])
        else:
            if in_table:
                if table_data: build_and_format_table(table_data)
                table_data, in_table = [], False

            if line_stripped.startswith("તા.") or line_stripped.startswith("સ્થળ:"):
                flush_signatures()
                p = doc.add_paragraph(line_stripped)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.paragraph_format.space_after = Pt(0) 
            elif "સાદર નોંધ" in line_stripped:
                flush_signatures()
                p = doc.add_paragraph()
                p.add_run(line_stripped).bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                pPr = p._p.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'dotted')
                bottom.set(qn('w:sz'), '4')
                bottom.set(qn('w:space'), '1')
                bottom.set(qn('w:color'), '000000')
                pBdr.append(bottom)
                pPr.append(pBdr)
            elif line_stripped.startswith("વિષય:"):
                flush_signatures()
                p = doc.add_paragraph()
                p.add_run(line_stripped).bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            elif len(line_stripped) < 60 and any(role in line_stripped for role in ["અધિકારી", "ઈન્ચાર્જ", "પ્રાધ્યાપક", "વડા"]) and not any(r in line_stripped for r in ["આચાર્ય", "ડીનશ્રી"]):
                sig_buffer.append(line_stripped)
            elif any(role in line_stripped for role in ["આચાર્ય", "ડીનશ્રી", "મહાવિધાયલય", "ન.કૃ.યુ"]):
                flush_signatures()
                doc.add_paragraph().paragraph_format.space_before = Pt(30)
                p_table = doc.add_table(rows=1, cols=2)
                p_table.columns[0].width, p_table.columns[1].width = Mm(79), Mm(74) 
                formatted_line = "\n".join([p.strip() for p in line_stripped.split(",")])
                p = p_table.cell(0, 1).paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run(formatted_line)
            else:
                flush_signatures()
                doc.add_paragraph(line_stripped).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY 

    flush_signatures()
    if in_table and table_data: build_and_format_table(table_data)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def create_purchase_order_docx(vendor_name, vendor_address, out_no, po_date, df_items):
    doc = Document()
    for section in doc.sections:
        section.top_margin, section.bottom_margin = Inches(0.4), Inches(0.5)
        section.left_margin, section.right_margin = Inches(0.8), Inches(0.8)
        
    style = doc.styles['Normal']
    style.font.size = Pt(12)
    style.paragraph_format.space_after, style.paragraph_format.space_before = Pt(0), Pt(0)
        
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    table.columns[0].width, table.columns[1].width, table.columns[2].width = Inches(1.8), Inches(3.6), Inches(1.4)
    
    if 'NAU_LOGO' in globals() and NAU_LOGO and os.path.exists(NAU_LOGO):
        cell_left = table.cell(0, 0)
        cell_left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p_left = cell_left.paragraphs[0]
        p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_left.add_run().add_picture(NAU_LOGO, width=Inches(1.8))
        
    p_center = table.cell(0, 1).paragraphs[0]
    p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_center.paragraph_format.line_spacing = 0.85
    r1 = p_center.add_run("કીટકશાસ્ત્ર વિભાગ\n")
    r1.bold = True
    r1.font.size = Pt(22)
    r2 = p_center.add_run("ન. મ. કૃષિ મહાવિદ્યાલય\nનવસારી કૃષિ યુનિવર્સિટી\nનવસારી- ૩૯૬ ૪૫૦ (ગુજરાત)")
    r2.bold, r2.font.size = True, Pt(14)
    
    if 'ICAR_LOGO' in globals() and ICAR_LOGO and os.path.exists(ICAR_LOGO):
        p_right = table.cell(0, 2).paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_right.add_run().add_picture(ICAR_LOGO, width=Inches(1.5))
        
    p_thick1 = doc.add_paragraph()
    p_thick1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p_thick1.paragraph_format.line_spacing = Pt(1)
    p_thick1.add_run().font.size = Pt(1) 
    add_bottom_border(p_thick1, size='24')
    
    table2 = doc.add_table(rows=1, cols=2)
    table2.columns[0].width, table2.columns[1].width = Inches(3.4), Inches(3.4)

    table2.cell(0,0).paragraphs[0].add_run("ડૉ. જે. જે. પસ્તાગીયા\nપ્રાધ્યાપક અને વડા")
    p2 = table2.cell(0,1).paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.add_run("મોબાઇલ: +૯૧ ૯૪૨૭૮ ૬૭૯૨૫\nઇમેલ: headentonau@gmail.com")
    
    for cell in table2.rows[0].cells:
        for p in cell.paragraphs: p.paragraph_format.line_spacing = 0.8
            
    p_thick2 = doc.add_paragraph()
    p_thick2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p_thick2.paragraph_format.line_spacing = Pt(1)
    p_thick2.add_run().font.size = Pt(1) 
    add_bottom_border(p_thick2, size='24')
    
    letter_year = po_date.split('/')[-1] if '/' in po_date else (po_date.split('.')[-1] if '.' in po_date else "૨૦૨૬")
    
    table3 = doc.add_table(rows=1, cols=2)
    table3.autofit = False
    
    # 6.9 Inches ની અંદર રહે તે રીતે પહોળાઈ સેટ કરી છે
    table3.columns[0].width = Inches(5.0)  
    table3.columns[1].width = Inches(1.9)  
    table3.cell(0, 0).width = Inches(5.0)
    table3.cell(0, 1).width = Inches(1.9)
    
    p_out = table3.cell(0,0).paragraphs[0]
    p_out.paragraph_format.space_before = Pt(0)
    p_out.paragraph_format.space_after = Pt(0)
    p_out.add_run(f"જા.નં. એસીએન/એન્ટો/એઆઈએનપી-એએ/{out_no}/{letter_year}, નવસારી")
    
    p_date = table3.cell(0,1).paragraphs[0]
    p_date.paragraph_format.space_before = Pt(0)
    p_date.paragraph_format.space_after = Pt(0)
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_date.add_run(f"તારીખ: {po_date}")
            
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    p_to = doc.add_paragraph()
    p_to.add_run("પ્રતિ,").bold = True
    doc.add_paragraph(vendor_name).runs[0].bold = True
    doc.add_paragraph(vendor_address)

    doc.add_paragraph()
    
    p_subj = doc.add_paragraph()
    p_subj.add_run("        વિષય: ખરીદી હુકમ").bold = True
    
    doc.add_paragraph("        જય ભારત સહ ઉપરોક્ત વિષય અન્વયે જણાવવાનું કે, અત્રેના કીટકશાસ્ત્ર વિભાગ ખાતે નિચેની વસ્તુઓ બિલ સહિત રજુ કરવા વિનંતી.").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph() 

    # --- વસ્તુઓના લિસ્ટ વાળા ટેબલની ગોઠવણ ---
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.autofit = False  # <--- આ ઉમેરવું ખૂબ જરૂરી છે
    
    # કુલ 6.9 Inches થવું જોઈએ (0.5 + 3.4 + 1.0 + 1.0 + 1.0 = 6.9)
    widths = [Inches(0.5), Inches(3.4), Inches(1.0), Inches(1.0), Inches(1.0)]
    
    headers = ["અ.નં.", "વસ્તુઓના નામ", "જથ્થો", "ભાવ પ્રતિ નંગ", "કુલ રકમ"]
    for i, ht in enumerate(headers):
        table.columns[i].width = widths[i]
        table.cell(0, i).width = widths[i] # હેડર સેલની પહોળાઈ લોક કરો
        table.cell(0, i).text = ht
        p = table.cell(0, i).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        table.cell(0, i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    total_amount = 0.0
    for index, row in df_items.iterrows():
        row_cells = table.add_row().cells
        
        # દરેક નવી રો ની પહોળાઈ લોક કરો
        for i in range(5): 
            row_cells[i].width = widths[i]

        row_cells[0].text = str(index + 1)
        row_cells[1].text = str(row.get('Details', ''))
        row_cells[2].text = str(row.get('Required Quantity', '')) + " " + str(row.get('Available Pkt/Unit', ''))
        
        unit_price = pd.to_numeric(row.get('Unit/Pkt Price', 0), errors='coerce')
        total_price = pd.to_numeric(row.get('Total Price', 0), errors='coerce')
        row_cells[3].text, row_cells[4].text = f"{unit_price:.2f}", f"{total_price:.2f}"
        total_amount += total_price
        
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    total_row = table.add_row().cells
    for i in range(5): 
        total_row[i].width = widths[i]

    total_row[3].text = "Total"
    total_row[3].paragraphs[0].runs[0].bold = True
    total_row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    total_row[4].text = f"{total_amount:.2f}"
    total_row[4].paragraphs[0].runs[0].bold = True
    total_row[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    p_sig = doc.add_paragraph()
    p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_sig.add_run("પ્રાધ્યાપક અને વડા\nકીટકશાસ્ત્ર વિભાગ").bold = True

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()
    
# --- Bill Payment Form ---
def create_bill_payment_form(budget_head, bill_no, bill_date, party_name, amount, amount_words):
    amount = coerce_amount(amount)
    doc = Document()
    for section in doc.sections:
        section.top_margin, section.bottom_margin = Inches(0.8), Inches(0.8)
        section.left_margin, section.right_margin = Inches(1.0), Inches(1.0)
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(15)
    style.paragraph_format.space_after = Pt(0)
    
    p_header = doc.add_paragraph()
    p_header.add_run("No. ACN/ENTO/BILL/       /202\n").bold = True
    r_right = p_header.add_run(f"NAVSARI-396450, Date: {datetime.date.today().strftime('%d/%m/%Y')}").bold = True
    p_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    
    doc.add_paragraph("To,").bold = True
    doc.add_paragraph("The Principal and Dean,").bold = True
    doc.add_paragraph("N.M. College of Agriculture,").bold = True
    doc.add_paragraph("Navsari").paragraph_format.space_after = Pt(12)
    
    p_sub = doc.add_paragraph()
    p_sub.add_run("Sub: Submission of bill(s) for payment............").bold = True
    p_sub.paragraph_format.space_after = Pt(12)
    
    p_body = doc.add_paragraph("With reference to the above subject, I am submitting herewith the following bill(s) for making payment to the respective party and debit the same in Budget Head No- ")
    p_body.add_run(budget_head).bold = True
    p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    
    hdr0 = table.rows[0].cells
    hdr1 = table.rows[1].cells
    
    hdr0[0].text = "Sr."
    hdr1[0].text = "No."
    hdr0[1].text = "No. of Bill/Date"
    hdr0[2].text = "Name of the Party"
    hdr0[3].text = "Amount"
    hdr1[3].text = "Rs.              Ps."
    
    for c in [0, 1, 2, 3]:
        if c in [1, 2]: table.cell(0, c).merge(table.cell(1, c))
        for r in [0, 1]:
            p = table.cell(r, c).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs: run.bold = True
            table.cell(r, c).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    table.columns[0].width = Inches(0.5)
    table.columns[1].width = Inches(1.5)
    table.columns[2].width = Inches(12.5)
    table.columns[3].width = Inches(1.5)

    row = table.add_row().cells
    row[0].text = "1"
    row[1].text = f"No: {bill_no}\nDt: {bill_date}"
    row[2].text = party_name
    row[3].text = format_amount(amount)
    
    for i in range(4): 
        row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    total_row = table.add_row().cells
    table.cell(3,0).merge(table.cell(3,2))
    p_tot = total_row[0].paragraphs[0]
    p_tot.add_run("Total:   ").bold = True
    p_tot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    total_row[3].text = format_amount(amount)
    total_row[3].paragraphs[0].runs[0].bold = True
    total_row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    word_row = table.add_row().cells
    table.cell(4,0).merge(table.cell(4,3))
    p_word = word_row[0].paragraphs[0]
    p_word.add_run("In words: ").bold = True
    p_word.add_run(f"Rupees {amount_words} Only.")
    
    party_row = table.add_row().cells
    table.cell(5,0).merge(table.cell(5,3))
    p_party = party_row[0].paragraphs[0]
    p_party.add_run("Name of Party for Payment: ").bold = True
    p_party.add_run(party_name)
         
    doc.add_paragraph("Encl: Cash/Credit Bill in original")
    doc.add_paragraph(f"No. {bill_no} with entry").paragraph_format.space_after = Pt(20)
    
    doc.add_paragraph(f"Copy F.W.C.S. to M/S: {party_name}").paragraph_format.space_after = Pt(30)
    
    p_sig = doc.add_paragraph()
    p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_sig.add_run("Professor and Head\nDepartment of Entomology\nN.M.C.A., N.A.U., Navsari").bold = True

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- PERFECT PDF REPLICA: Bill Pasting Form ---
def set_cell_border(cell, **kwargs):
    """Helper function to draw borders around specific table cells."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

import io
from docx import Document
from docx.shared import Inches, Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    """Helper function to draw borders around specific table cells."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def create_bill_pasting_form(budget_head, grant_year, party_name, amount, amount_in_guj_words, reg_type, reg_page_no, bill_reg_date, bill_reg_page_no, bill_reg_sr_no, item_no, approval_no, approval_date):
    amount = coerce_amount(amount)
    doc = Document()
    
    # --- Helper to convert English digits to Gujarati digits ---
    def eng_to_guj(text):
        if not text: return ""
        return str(text).translate(str.maketrans("0123456789", "૦૧૨૩૪૫૬૭૮૯"))

    # Strict A4 Margins
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.2)
    section.bottom_margin = Inches(0.2)
    section.gutter = Inches(0)
    
    # Set base fonts
    style = doc.styles['Normal']
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Shruti')
    style.font._element.append(rFonts)

    # --- PAGE 1 ---
    header_table = doc.add_table(rows=1, cols=3)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.columns[0].width = Inches(2.0)
    header_table.columns[1].width = Inches(2.7) 
    header_table.columns[2].width = Inches(2.0)
    
    cell_left = header_table.cell(0, 0)
    p_office = cell_left.paragraphs[0]
    p_office.paragraph_format.space_before = Pt(12)
    run_office = p_office.add_run("Office No. 303")
    run_office.bold = True
    run_office.font.size = Pt(20) 
    p_office.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_border(cell_left, top={"sz": 12, "val": "single", "color": "000000"}, bottom={"sz": 12, "val": "single", "color": "000000"}, left={"sz": 12, "val": "single", "color": "000000"}, right={"sz": 12, "val": "single", "color": "000000"})
    
    cell_right = header_table.cell(0, 2)
    p_v1 = cell_right.paragraphs[0]
    run_v1 = p_v1.add_run("Voucher No:-....................")
    run_v1.font.size = Pt(15) 
    p_v1.paragraph_format.space_before = Pt(6) 
    p_v1.paragraph_format.space_after = Pt(8)  
    
    p_v2 = cell_right.add_paragraph()
    run_v2 = p_v2.add_run("          Date:-....................")
    run_v2.font.size = Pt(15) 
    p_v2.paragraph_format.space_after = Pt(0)
    set_cell_border(cell_right, top={"sz": 12, "val": "single", "color": "000000"}, bottom={"sz": 12, "val": "single", "color": "000000"}, left={"sz": 12, "val": "single", "color": "000000"}, right={"sz": 12, "val": "single", "color": "000000"})
        
    p_col = doc.add_paragraph()
    p_col.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_col.paragraph_format.space_before = Pt(0) 
    p_col.paragraph_format.space_after = Pt(0)
    run_col = p_col.add_run("N. M. COLLEGE OF AGRICULTURE")
    run_col.bold = True
    run_col.font.size = Pt(22) 
    
    p_uni = doc.add_paragraph()
    p_uni.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_uni = p_uni.add_run("Navsari Agricultural University, Navsari.-396450")
    run_uni.bold = True
    run_uni.font.size = Pt(20) 
    p_uni.paragraph_format.space_after = Pt(10)
    
    line_table = doc.add_table(rows=1, cols=1)
    line_table.columns[0].width = Inches(6.7)
    set_cell_border(line_table.cell(0, 0), top={"sz": 24, "val": "single", "color": "000000"})
    
    for _ in range(15): doc.add_paragraph()

    line_table_2 = doc.add_table(rows=1, cols=1)
    line_table_2.columns[0].width = Inches(6.7)
    set_cell_border(line_table_2.cell(0, 0), top={"sz": 24, "val": "single", "color": "000000"})
    
    p_note = doc.add_paragraph()
    run_note = p_note.add_run("Note:-")
    run_note.bold = True
    run_note.font.size = Pt(15) 
    p_note.paragraph_format.space_after = Pt(6)
    
    def add_bullet(num, text, size=15):
        p = doc.add_paragraph()
        run_num = p.add_run(num + "\t")
        run_num.font.size = Pt(15) 
        run_text = p.add_run(text)
        run_text.font.size = Pt(15) 
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(2)
        
    add_bullet("1.", "Quotation from at least three parties for purchase above Rs. 1000/- should be obtained.", size=10)
    add_bullet("2.", "Purchase from authorized details and manufactures be certified on bill no other quotation were available due of purchase from manufacture or authorized dealers.", size=10)
    add_bullet("3.", "A special previous sanction of V.C of campus, Navsari should invariably be obtained for dead stock of other valuable articles before the purchase is made.", size=10)
    add_bullet("4.", "Purchase is made in the interested of University work.", size=10)
    
    doc.add_page_break()

    # --- PAGE 2 ---
    def add_p2_header_row(cell, text, size=10):
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0) 
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.font.size = Pt(size)
        
    top_table = doc.add_table(rows=4, cols=3)
    top_table.autofit = False
    top_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
    for row in top_table.rows:
        row.cells[0].width = Inches(2.0)
        row.cells[1].width = Inches(0.3)
        row.cells[2].width = Inches(4.9)

    add_p2_header_row(top_table.cell(0,0), "બજેટ સદર")
    add_p2_header_row(top_table.cell(0,1), ":-")
    
    p_0_2 = top_table.cell(0,2).paragraphs[0]
    p_0_2.paragraph_format.space_before = Pt(0)
    p_0_2.paragraph_format.space_after = Pt(0)
    p_0_2.paragraph_format.left_indent = Inches(0)
    run_bh = p_0_2.add_run(f"{budget_head} \t\t")
    run_bh.font.size = Pt(10)
    run_exp = p_0_2.add_run("EXP. CODE NO._________")
    run_exp.font.size = Pt(10)
    run_exp.bold = True

    add_p2_header_row(top_table.cell(1,0), "ફાળવેલ ગ્રાન્ટ વર્ષ: ૨૦  - ૨૦")
    add_p2_header_row(top_table.cell(1,1), ":-")
    add_p2_header_row(top_table.cell(1,2), f"{grant_year}" if grant_year else "                     ")
    
    add_p2_header_row(top_table.cell(2,0), "બીલની કુલ રકમ")
    add_p2_header_row(top_table.cell(2,1), ":-")
    add_p2_header_row(top_table.cell(2,2), format_amount(amount)) 
    
    add_p2_header_row(top_table.cell(3,0), "ચુકવણું કરવામાં આવનાર પાર્ટીનું નામ (અંગ્રેજી કેપીટલ લેટર)")
    add_p2_header_row(top_table.cell(3,1), ":-")
    add_p2_header_row(top_table.cell(3,2), f"{party_name}")
    
    p_cert = doc.add_paragraph()
    p_cert.paragraph_format.space_before = Pt(6) 
    p_cert.paragraph_format.space_after = Pt(4)  
    run_cert = p_cert.add_run(":: પ્રમાણપત્ર ::")
    run_cert.bold = True
    run_cert.font.size = Pt(14)
    p_cert.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    table = doc.add_table(rows=9, cols=2)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
    for row in table.rows:
        row.cells[0].width = Inches(0.4)
        row.cells[1].width = Inches(6.8)
    
    def add_row(idx, no, text, size=11):
        p_no = table.rows[idx].cells[0].paragraphs[0]
        p_no.paragraph_format.left_indent = Inches(0) 
        p_no.paragraph_format.space_before = Pt(0)
        p_no.paragraph_format.space_after = Pt(0)
        run_no = p_no.add_run(no)
        run_no.font.size = Pt(9)
        
        p_text = table.rows[idx].cells[1].paragraphs[0]
        p_text.paragraph_format.left_indent = Inches(0)
        p_text.paragraph_format.space_before = Pt(0)
        p_text.paragraph_format.space_after = Pt(2) 
        p_text.paragraph_format.line_spacing = 1.0  
        p_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # નોંધ: અહીથી જૂનો `run_text = p_text.add_run(text)` વાળો કોડ કાઢી નાખ્યો છે.
        
        # --- NEW: ** વચ્ચેના લખાણને Bold કરવા માટેનું સ્માર્ટ લોજીક ---
        parts = str(text).split("**")
        for i, part in enumerate(parts):
            if part:
                run_text = p_text.add_run(part)
                run_text.font.size = Pt(9)
                # એકી સંખ્યા વાળા ભાગ (અર્થાત ** ની વચ્ચે વાળા) Bold થશે
                if i % 2 != 0:  
                    run_text.bold = True

    # --- Convert inputs to Gujarati digits ---
    guj_amount = eng_to_guj(f"{amount:.2f}")
    guj_reg_page = eng_to_guj(reg_page_no)
    guj_bill_reg_page = eng_to_guj(bill_reg_page_no)
    guj_bill_reg_sr = eng_to_guj(bill_reg_sr_no)
    guj_bill_date = eng_to_guj(bill_reg_date)
    
    # New: Convert Approval Details to Gujarati
    guj_item_no = eng_to_guj(item_no) if item_no else "____________"
    guj_app_no = eng_to_guj(approval_no) if approval_no else "_________________________________________"
    guj_app_date = eng_to_guj(approval_date) if approval_date else "______/______/_________"

    # --- Register Setup ---
    blanks = {
        "સ્ટોર રોજમેળ": "____________",
        "ચીજવસ્તુ વપરાશ (કન્ઝયુમેબલ)": "____________",
        "ડેડસ્ટોક": "____________",
        "ટેલીફોન": "____________",
        "સ્ટેમ્પ": "____________",
        "સ્ટેશનરી": "____________",
        "પરચુરણ માલ સામાન": "____________",
        "રીપેરીંગ": "____________"
    }

    if guj_reg_page:
        if reg_type in blanks:
            blanks[reg_type] = f" {guj_reg_page} "

    # --- Dynamic Certificates ---
    cert_1 = (f"આ બીલમાં જણાવેલ વસ્તુ ખરીદવાની/રીપેરીંગના ખર્ચની મંજુરી આપવાની સત્તા ગુજરાત રાજય કૃષિ યુનિવર્સિટીઓ "
              f"(સતા સોપણી) નિયમ-૨૦૧૧ ના સ્ટેચ્યુટ નં. ૧૨૧ ની આઇટમ નં {guj_item_no} મુજબ એનાયત થયેલ સત્તા પ્રમાણે "
              f"હેડ ઓફિસ/હેડ ઓફ યુનિટ/યુનિ. ઓફિસર્સ/માન. કુલપતિશ્રીની મંજુરી નં: {guj_app_no} . "
              f"તારીખ: {guj_app_date} થી મંજુરી મળેલ છે. હુકમની નકલ સામેલ છે.")

    cert_2 = "આ બીલમાં જણાવેલ ખર્ચ આ વિભાગની આઇ.સી.એ.આર. યોજના બજેટ સદર ૩૦૩/૨૦૯૨ માં સમાવેશ કરવામાં આવેલ છે."
    
    cert_3 = (
        f"બીલમાં દર્શાવેલ માલની ખરીદી બજાર ભાવ તપાસી ભાવો મેળવી સૌથી ઓછા ભાવ મુજબ છે અને સારી સ્થિતિમાં મળેલ છે. "
        f"જે કચેરીના સ્ટોર રોજમેળ રજી પાના નં. {blanks['સ્ટોર રોજમેળ']}../ચીજવસ્તુ વપરાશ (કન્ઝયુમેબલ) "
        f"રજી. પાના નં. {blanks['ચીજવસ્તુ વપરાશ (કન્ઝયુમેબલ)']} ડેડસ્ટોક રજી. નં.... {blanks['ડેડસ્ટોક']} / ટેલીફોન રજી. પાના "
        f"નં {blanks['ટેલીફોન']} / સ્ટેમ્પ રજી. પાના નં {blanks['સ્ટેમ્પ']} / સ્ટેશનરી રજી. પાના નં. "
        f"{blanks['સ્ટેશનરી']} .રજીસ્ટરનાં ____________ / પરચુરણ માલ સામાન /.... {blanks['પરચુરણ માલ સામાન']}../ રીપેરીંગ "
        f"રજી. પાના નં.. {blanks['રીપેરીંગ']} નાં રોજ જમા કરવામાં આવેલ છે."
    )
    
    cert_4 = f"સદર બીલમાં દર્શાવવામાં આવેલ ખર્ચ સેલ્સ ટેક્ષ / એડી.ટેક્ષ / એકસાઇડયુટી / સેન્ટ્રલ ટેક્ષ વિગેરે પાર્ટીના માન્ય થયેલ ભાવ મુજબ ચકાસણી કરવામાં આવેલ છે. અને તે મુજબ પાર્ટીના બીલમાં દર્શાવ્યા મુજબની રકમ રૂ. {guj_amount}/- (અંકે રૂ. {amount_in_guj_words}) પુરા ચુકવવા ભલામણ કરવામાં આવે છે."
    cert_8 = f"તા. {guj_bill_date} સદર બીલની નોંધ કચેરી ખાતેના બીલ રજી પાના નં. {guj_bill_reg_page} અનુ.નં. {guj_bill_reg_sr} કરવામાં આવેલ છે."

    add_row(0, "૧.", cert_1, size=11)
    add_row(1, "૨.", cert_2, size=11)
    add_row(2, "૩.", cert_3, size=11)
    add_row(3, "૪.", cert_4, size=11)
    add_row(4, "૫.", "બીલમાં દર્શાવેલ મુજબ વાહન નં.... ____________ .ની રીપેરીંગ કામગીરી સંતોષકારક થયેલ છે જેની નોંધ હિસ્ટ્રીશીટ રજી. પાના નં....... ____________ .../ રીપેરીંગ રજી. પાના નં. ____________ થી કરેલ છે. જે ચાલુ નાણાંકીય વર્ષ દરમ્યાન આ બીલ સહીત કુલ ખર્ચ રૂ... ____________./- (અંકે ૩....................................) થયેલ છે.", size=11)
    add_row(5, "૬.", "બીલમાં દર્શાવેલ પેટ્રોલ / ડીઝલ / ઓઇલ વગેરે વાહન નં. ____________ .માટે ખરીદ કરવામાં આવેલ છે જે લોગબુક ભાગ નં. ____________ પાના નં. ____________ થી જમાં કરવામાં આવેલ છે.", size=11)
    add_row(6, "૭.", "તા. ____________ બીલમાં દર્શાવેલ વાહન નં. ____________ ના રીપેરીંગ કામ કરતી વખતે પરત આવેલ જુના સ્પેર પાર્ટસ મેળવીને આ કચેરીનાં રદ્દ રજી. પાના નં. ____________ ના રોજ જમાં લીધેલ છે.", size=11)
    add_row(7, "૮.", cert_8, size=11)
    add_row(8, "૯.", "સંશોધન નિયામકશ્રીના ૩૦/૧૦/૨૦૨૧ના પરિપત્રનો અમલ કરેલ છે.", size=11)
    
    p_special_note = doc.add_paragraph()
    p_special_note.paragraph_format.space_before = Pt(0)
    p_special_note.paragraph_format.space_after = Pt(2) 
    p_special_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_special_note = p_special_note.add_run("સદરહું ખર્ચ કચેરીની અગત્યની કામગીરીને ધ્યાને લઇ તેમજ યુનિવર્સિટીનાં હિતાર્થે કરવામાં આવેલ છે.")
    run_special_note.font.size = Pt(11)
    
    today_guj = eng_to_guj(datetime.date.today().strftime('%d/%m/%Y'))
    p_loc = doc.add_paragraph()
    p_loc.paragraph_format.space_before = Pt(2) 
    p_loc.paragraph_format.space_after = Pt(12)
    run_loc = p_loc.add_run(f"સ્થળ : નવસારી\nતારીખ : {today_guj}")
    run_loc.font.size = Pt(12)
    
    # --- સુધારેલો ભાગ: ૨ કોલમને બદલે ૩ સહીના કોલમ (3 Signature Columns) ---
    table_sig = doc.add_table(rows=1, cols=3)
    table_sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # કુલ લિમિટ (6.8 Inches) ને ૩ સરખા ભાગમાં વહેંચી દીધી (આશરે 2.26 ઇંચ દરેક)
    col_widths = [Inches(2.26), Inches(2.26), Inches(2.26)]
    for i, cell in enumerate(table_sig.rows[0].cells): 
        cell.width = col_widths[i]
    
    # કોલમ ૧: ખેતીવાડી અધિકારી (Left Alignment)
    p_s1 = table_sig.cell(0, 0).paragraphs[0]
    p_s1.paragraph_format.space_before, p_s1.paragraph_format.space_after = Pt(0), Pt(0)
    p_s1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_s1 = p_s1.add_run("ખેતીવાડી અધિકારીની\nસહી અને હોદ્દો")
    run_s1.bold, run_s1.font.size = True, Pt(11)
    
    # કોલમ ૨: સિનિયર અકેરોલોજીસ્ટ (Center Alignment)
    p_s2 = table_sig.cell(0, 1).paragraphs[0]
    p_s2.paragraph_format.space_before, p_s2.paragraph_format.space_after = Pt(0), Pt(0)
    p_s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_s2 = p_s2.add_run("સિનિયર અકેરોલોજીસ્ટની\nસહી અને હોદ્દો")
    run_s2.bold, run_s2.font.size = True, Pt(11)
    
    # કોલમ ૩: વિભાગીય વડા (Right Alignment)
    p_s3 = table_sig.cell(0, 2).paragraphs[0]
    p_s3.paragraph_format.space_before, p_s3.paragraph_format.space_after = Pt(0), Pt(0)
    p_s3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_s3 = p_s3.add_run("વિભાગીય વડાની\nસહી અને હોદ્દો")
    run_s3.bold, run_s3.font.size = True, Pt(11)
    
    # --- નીચેનો બાકીનો કોડ જેમ છે તેમ જ રહેશે ---
    p_passed = doc.add_paragraph()
    p_passed.paragraph_format.space_before, p_passed.paragraph_format.space_after = Pt(20), Pt(6)
    run_passed = p_passed.add_run("Passed for Payment Rs ........................................\nRupees: ........................................................................................")
    run_passed.font.size = Pt(12)
    
    p_aao = doc.add_paragraph()
    p_aao.paragraph_format.space_before, p_aao.paragraph_format.space_after = Pt(0), Pt(0)
    run_aao = p_aao.add_run("Assistant Administrative Officer\nN. M. College of Agriculture\nNavsari-396 450")
    run_aao.bold, run_aao.font.size = True, Pt(12)
    p_aao.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ==========================================
# Streamlit App UI
# ==========================================
st.set_page_config(page_title="સાદર નોંધ જનરેટર", layout="wide")
st.title("સાદર નોંધ જનરેટર (Intelligent Sadar Nondh App)")

st.session_state["_ui_ready"] = True
for warning_message in st.session_state.pop("_startup_warnings", []):
    st.sidebar.warning(warning_message)

api_key = get_secret_value(
    "GEMINI_API_KEY",
    "GOOGLE_API_KEY",
    "GOOGLE_GENERATIVE_AI_API_KEY",
)

if api_key:
    st.sidebar.success("Gemini API key loaded from Streamlit secrets.")
else:
    st.sidebar.warning('Gemini API key missing. Add GEMINI_API_KEY in Streamlit secrets.')

# --- ADDED TAB 6 ---
tab0, tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
    "Dashboard",
    "નવી સાદર નોંધ (Create)", 
    "જુની નોંધ (Archives)", 
    "ખરીદી હુકમ (Purchase Order)",
    "બિલ પેમેન્ટ (Bill Payment)",
    "બિલ પેસ્ટિંગ (Bill Pasting)",
    "🗄️ ડિજિટલ આર્કાઇવ (Digital Vault)"  # <-- NEW TAB
])

with tab0:
    st.markdown("### Workflow Dashboard")

    current_year_for_dashboard = datetime.date.today().year
    dashboard_fy_options = ["All"] + [f"{y}-{str(y + 1)[2:]}" for y in range(current_year_for_dashboard - 2, current_year_for_dashboard + 3)][::-1]
    dashboard_month_options = ["All"] + DASHBOARD_MONTHS

    filter_col1, filter_col2 = st.columns(2)
    with filter_col1:
        dashboard_fy = st.selectbox("Financial Year", dashboard_fy_options, key="dashboard_fy_filter")
    with filter_col2:
        dashboard_month = st.selectbox("Month", dashboard_month_options, key="dashboard_month_filter")

    dashboard_data = build_dashboard_data(dashboard_fy, dashboard_month)
    dashboard_metrics = dashboard_data["metrics"]

    if dashboard_data.get("errors"):
        with st.expander("Dashboard data warnings", expanded=False):
            for dashboard_error in dashboard_data["errors"]:
                st.warning(dashboard_error)

    metric_cards = [
        ("Total Nondhs", dashboard_metrics["total_nondhs"]),
        ("Awaiting PO", dashboard_metrics["nondhs_awaiting_po"]),
        ("Unfinished POs", dashboard_metrics["unfinished_pos"]),
        ("Payment Forms", dashboard_metrics["payment_forms_generated"]),
        ("Paid Bills", dashboard_metrics["paid_bills"]),
        ("Pending Amount", f"Rs. {format_amount(dashboard_metrics['pending_amount'])}"),
        ("Paid Amount", f"Rs. {format_amount(dashboard_metrics['paid_amount'])}"),
        ("Missing Documents", dashboard_metrics["missing_documents"]),
        ("Memory Queue", dashboard_metrics["pending_memory_suggestions"]),
        ("Skill Queue", dashboard_metrics["pending_skill_suggestions"]),
    ]

    for start in (0, 5):
        metric_cols = st.columns(5)
        for col, (label, value) in zip(metric_cols, metric_cards[start:start + 5]):
            with col:
                st.metric(label, value)

    if (
        dashboard_metrics["total_nondhs"] == 0
        and dashboard_metrics["unfinished_pos"] == 0
        and dashboard_metrics["payment_forms_generated"] == 0
        and dashboard_metrics["paid_bills"] == 0
    ):
        st.info("No workflow records found for the selected filters yet.")

    st.markdown("---")
    summary_col1, summary_col2 = st.columns(2)
    with summary_col1:
        st.markdown("#### Money Summary")
        st.dataframe(pd.DataFrame(dashboard_data["money_summary"]), use_container_width=True, hide_index=True)
    with summary_col2:
        st.markdown("#### Learning Queue")
        st.dataframe(pd.DataFrame(dashboard_data["learning_queue"]), use_container_width=True, hide_index=True)

    st.markdown("#### Next Actions")
    if dashboard_data["next_actions"]:
        st.dataframe(pd.DataFrame(dashboard_data["next_actions"]), use_container_width=True, hide_index=True)
    else:
        st.success("No urgent next actions for the selected filters.")

    st.markdown("#### Document Gaps")
    document_gaps = [row for row in dashboard_data["document_rows"] if row.get("Missing Count", 0) > 0]
    if document_gaps:
        st.warning(f"{sum(row.get('Missing Count', 0) for row in document_gaps)} required document slots are still missing.")
        st.dataframe(pd.DataFrame(document_gaps), use_container_width=True, hide_index=True)
    elif dashboard_data["document_rows"]:
        st.success("All tracked workflow document slots are complete for the selected filters.")
    else:
        st.info("No Nondh/PO workflow rows to check yet.")

    st.markdown("#### Pending Payments")
    if dashboard_data["pending_payments"]:
        st.dataframe(pd.DataFrame(dashboard_data["pending_payments"]), use_container_width=True, hide_index=True)
    else:
        st.success("No pending payment rows for the selected filters.")

    with st.expander("AI Workflow Coach", expanded=False):
        st.caption("Runs only when clicked and uses the dashboard summary plus approved learning memories/skills.")
        if st.button("Summarize Bottlenecks", key="dashboard_ai_workflow_coach"):
            if not api_key:
                st.warning("Add GEMINI_API_KEY in Streamlit secrets to use the AI Workflow Coach.")
            else:
                with st.spinner("Reviewing workflow bottlenecks..."):
                    try:
                        coach_payload = build_dashboard_coach_prompt(dashboard_data)
                        coach_learning = get_learning_context(
                            coach_payload,
                            ["general_workflow", "bill_payment", "bill_pasting", "invoice_extraction", "register_rule"],
                            "dashboard_coach",
                            memory_limit=6,
                            skill_limit=3,
                        )
                        genai.configure(api_key=api_key)
                        coach_model = genai.GenerativeModel('gemini-3.1-flash-lite-preview')
                        coach_prompt = f"""
You are the AI Workflow Coach for Nodh Maker.
Use only the dashboard JSON and approved learning context below. Do not request, reveal, store, or infer any API keys or secrets.
Return a concise operational summary with:
1. Biggest bottleneck.
2. Next 5 actions in priority order.
3. Document gaps to fix first.
4. Money/payment risk summary.
5. Learning queue recommendation.

Dashboard JSON:
{coach_payload}

Approved memories and Hermes-style skills:
{coach_learning['prompt'] if coach_learning['prompt'] else 'No approved learning context matched.'}
"""
                        coach_response = coach_model.generate_content(coach_prompt)
                        st.markdown(redact_sensitive(coach_response.text))
                        with st.expander("Coach context used", expanded=False):
                            if coach_learning["memories"]:
                                st.markdown("**Memories**")
                                for memory in coach_learning["memories"]:
                                    st.caption(f"#{memory['id']} {memory['title']} ({memory['category']})")
                            if coach_learning["skills"]:
                                st.markdown("**Skills**")
                                for skill in coach_learning["skills"]:
                                    st.caption(f"#{skill['id']} {skill['name']} v{skill['version']}")
                            if not coach_learning["memories"] and not coach_learning["skills"]:
                                st.caption("No approved memories or skills were selected.")
                    except Exception as e:
                        st.warning(f"AI Workflow Coach failed: {e}")

with tab1:
    st.markdown("### જરૂરિયાતની વિગત આપો (Provide Requirements)")

    nondh_type = st.radio(
        "નોંધનો પ્રકાર (Type of Nondh):",
        ["સામાન્ય ખરીદી (Standard Purchase)", "કાયમી પેશગી / એડવાન્સ (Advance Payment)"],
        horizontal=True,
        help="એવી સંસ્થા/વેન્ડર માટે કે જે કામ પૂર્ણ થાય તે પહેલાં જ પેમેન્ટ માંગે (દા.ત. GBRC) ત્યાં 'કાયમી પેશગી / એડવાન્સ' પસંદ કરો.",
    )
    is_advance = "પેશગી" in nondh_type

    advance_vendor = ""
    if is_advance:
        advance_vendor = st.text_input(
            "ક્વોટેશન આપનાર સંસ્થા/વેન્ડરનું નામ (Institute providing quotation):",
            value="ગુજરાત બાયોટેકનોલોજી રિસર્ચ સેન્ટર (GBRC), ગાંધીનગર",
            help="જે સંસ્થા સેવા આપશે અને એડવાન્સ પેમેન્ટ માંગે છે તેનું નામ.",
        )

    col1, col2 = st.columns(2)
    with col1:
        text_prompt = st.text_area("તમારી જરૂરિયાત લખો:", placeholder="e.g., need 10 entomological pins...")
    with col2:
        uploaded_image = st.file_uploader("અથવા PDF/ફોટો અપલોડ કરો:", type=["pdf", "jpg", "jpeg", "png"])

    if st.button("જનરેટ કરો (Generate)"):
        if not api_key:
            st.error("Please add GEMINI_API_KEY in Streamlit secrets.")
        elif not text_prompt and not uploaded_image:
            st.warning("Please provide either a text requirement or a PDF/image.")
        else:
            with st.spinner("સ્ટેચ્યુટ ૧૨૧ ની ચકાસણી અને નોંધ તૈયાર કરવામાં આવી રહી છે..."):
                try:
                    statute_context, sample_context = load_permanent_context()
                    learning = get_learning_context(
                        text_prompt or "",
                        ["nondh_style", "statute_precedent", "item_mapping", "budget_default", "general_workflow"],
                        "nondh_generation"
                    )

                    genai.configure(api_key=api_key)
                    model = genai.GenerativeModel('gemini-3.1-flash-lite-preview')

                    common_context = f"""
                    [CONTEXT START]
                    {statute_context[:15000]}
                    {sample_context}
                    [CONTEXT END]

                    [APPROVED LEARNING MEMORY AND HERMES-STYLE SKILLS]
                    {learning['prompt'] if learning['prompt'] else 'No approved memories or skills matched this request.'}
                    [END APPROVED LEARNING]
                    """

                    common_table_rules = """
                    If a table is included, the headers MUST strictly be in ENGLISH and use these EXACT columns:
                    Sr. No. | Details | Required Quantity | Available Pkt/Unit | Unit/Pkt Price | Total Price

                    CRITICAL TABLE RULES:
                    1. The 'Details' column MUST contain ONLY the item/service name without the package size (e.g., "Sanger Sequencing- COI Gene Sequencing").
                    2. The 'Available Pkt/Unit' column MUST contain the package size and unit type (e.g., "500 ML", "25 GM", "1 Unit").
                    3. The 'Required Quantity' and 'Unit/Pkt Price' columns MUST contain pure numbers only (e.g., "1" or "1416"). Do NOT put units like "ML" or "GM" in these two columns.
                    4. Do NOT generate a "Grand Total" row. The system will calculate it automatically.
                    """

                    common_signatures = """
                    ખેતીવાડી અધિકારી,કીટકશાસ્ત્ર વિભાગ
                    પ્રોજેકટ ઈન્ચાર્જ,કીટકશાસ્ત્ર વિભાગ
                    પ્રાધ્યાપક અને વડા,કીટકશાસ્ત્ર વિભાગ

                    આચાર્ય અને ડીનશ્રી, ન. મ. કૃષિ મહાવિધાયલય, ન.કૃ.યુ. નવસારી
                    """

                    if is_advance:
                        sys_prompt = f"""
                    You are an expert administrative AI for the Department of Entomology, N. M. College of Agriculture, NAU, Navsari.
                    Your task is to generate a formal 'સાદર નોંધ' in Gujarati of the type "કાયમી પેશગી / એડવાન્સ પેશગી" (ADVANCE PAYMENT note).

                    WHEN THIS TYPE IS USED: An outside institute/vendor (here: {advance_vendor or 'ગુજરાત બાયોટેકનોલોજી રિસર્ચ સેન્ટર (GBRC), ગાંધીનગર'}) provides a service (e.g. Sanger Sequencing / COI Gene, molecular identification, testing) and DEMANDS PAYMENT IN ADVANCE, before the work is completed. Therefore the department must first withdraw the money as an advance (પેશગી) and pay the institute itself. This note seeks the Dean's in-principle (સૈદ્ધાંતિક) sanction to withdraw that advance.

                    {common_context}

                    Format REQUIRED (follow this wording style closely, filling the brackets from the user's request):
                    તા. {datetime.date.today().strftime('%d/%m/%Y')}
                    સ્થળ: નવસારી
                    સાદર નોંધ:
                    વિષય: કાયમી પેશગી પેટે [Service / work name, e.g. Mite Species નું Molecular Identification] માટે નાણા ઉપાડવા બાબત...
                    સવિનય સહ ઉપરોક્ત વિષય અન્વયે જણાવવાનું કે, અત્રેનાં કીટકશાસ્ત્ર વિભાગની આઈ.સી.એ.આર. યોજના AINP on Agril Acarology બ.સ. ૩૦૩/૨૦૯૨ અંતર્ગત [detailed logical reason describing the research need and why the outside service is required]. આ માટે {advance_vendor or 'ગુજરાત બાયોટેકનોલોજી રિસર્ચ સેન્ટર (GBRC), ગાંધીનગર'} દ્વારા મળેલ ક્વોટેશન મુજબ અંદાજિત ખર્ચ [Total Amount]/- (અંકે રૂપિયા [amount in Gujarati words] પૂરા) થનાર છે. જેના માટે [Total Amount]/- (અંકે રૂપિયા [amount in Gujarati words] પૂરા) એડવાન્સ પેટે ઉપાડવાની જરૂરિયાત ઉપસ્થિત થયેલ છે.

                    [INSERT THE MARKDOWN TABLE HERE]

                    સદર ખર્ચની ખરીદી કરવા આપશ્રીની સત્તા અન્વયે કાયમી પેશગી ઉપાડવાની સૈદ્ધાંતિક મંજુરી આપવા આપ સાહેબશ્રીને નમ્ર વિનંતી છે. સદર નાણા બ.સ ૩૦૩/૨૦૯૨ માંથી ફાળવી સદર યોજનાના પ્રોજેક્ટ ઈન્ચાર્જશ્રીને સોંપવા નમ્ર વિનંતી.

                    ADVANCE-PAYMENT SPECIFIC RULES:
                    1. This is an advance (પેશગી) note, so DO NOT ask for a specific "સ્ટેચ્યુટ ૧૨૧ આઈટમ નંબર" in the main body. The sanction is granted under the Dean's authority to release advance/પેશગી from the scheme funds.
                    2. Always name the quotation-giving institute exactly as: {advance_vendor or 'ગુજરાત બાયોટેકનોલોજી રિસર્ચ સેન્ટર (GBRC), ગાંધીનગર'}.
                    3. State the amount identically in both the "અંદાજિત ખર્ચ" sentence and the "એડવાન્સ પેટે ઉપાડવાની" sentence, with the Gujarati words in brackets.
                    4. Always end with the request to release the funds from બ.સ ૩૦૩/૨૦૯૨ and hand them to the Project In-charge.

                    TABLE LOGIC:
                    An advance-payment note ALWAYS lists the service(s) being paid for, so YOU MUST include a markdown table.
                    {common_table_rules}
                    {common_signatures}
                    ==== AI STATUTE ANALYSIS ====
                    1. **Nature of Sanction:** Explain that this is an advance/કાયમી પેશગી released under the Dean's financial powers because {advance_vendor or 'the institute'} demands payment before delivering the service.
                    2. **Justification:** Explain why paying this institute in advance is necessary for the research (relate to the AINP Acarology scheme).
                    3. **Similar Precedent from Sample Nondh:** If any advance/પેશગી precedent exists in the Sample Nondh context, cite its Subject, Date and wording. If none, state that this follows the standard GBRC advance-payment format.
                    4. **Fund Source:** Confirm the funds are drawn from AINP on Agril Acarology (બ.સ. ૩૦૩/૨૦૯૨).
                    """
                    else:
                        sys_prompt = f"""
                    You are an expert administrative AI for the Department of Entomology, N. M. College of Agriculture, NAU, Navsari.
                    Your task is to generate a formal 'સાદર નોંધ' in Gujarati.

                    {common_context}

                    Format REQUIRED:
                    તા. {datetime.date.today().strftime('%d/%m/%Y')}
                    સ્થળ: નવસારી
                    સાદર નોંધ:
                    વિષય: [Appropriate Subject...]
                    સવિનય ઉપરોક્ત વિષય અન્વયે જણાવવાનું કે, અત્રેનાં કીટકશાસ્ત્ર વિભાગની આઈ.સી.એ.આર. યોજના AINP on Agril Acarology બ.સ. ૩૦૩/૨૦૯૨ અંતર્ગત [Detailed logical reason]. સદર વસ્તુનો કુલ અંદાજિત ખર્ચ [Total Amount] થનાર છે.
                    જે આપ સાહેબશ્રીને સ્ટેચ્યુટ ૧૨૧ની આઈટમ નંબર [DETERMINED_ITEM_NUMBER] મુજબ એનાયત થયેલ સત્તા અનુસાર સૈદ્ધાંતિક મંજુરી આપવા વિનંતી. સદર ખર્ચ અત્રેના વિભાગમાં ચાલતી આઈ.સી.એ.આર યોજના (બ.સ. ૩૦૩/૨૦૯૨) માં કરવામાં આવશે.

                    STATUTE 121 ITEM NUMBER DETERMINATION (CRITICAL INSTRUCTION):
                    1. You MUST read the 'Sample Nondh Format' provided in the context to find the historically correct 'આઈટમ નંબર' (Item Number) for this type of purchase.
                    2. Treat the Sample Nondh as the ultimate precedent. If the purchase involves laboratory chemicals, research materials, or AINP scheme-related items, strictly use the exact item number found in the sample (e.g., "૫૪ (i)" or "54 (i)").
                    3. Only rely on the dense 'Statute 121 Rules' PDF text if the item category is completely new and not covered by the sample document precedent.
                    4. Replace [DETERMINED_ITEM_NUMBER] with the exact number in Gujarati format (like ૫૪ (i)). Do not guess or hallucinate numbers like 45 (ii) (ii).

                    TABLE LOGIC:
                    Analyze the user request to determine if a table is required.
                    - If the request is a general administrative note WITHOUT specific items to purchase, DO NOT include a table.
                    - If the request involves purchasing, requesting, or listing items with quantities and prices, YOU MUST include a markdown table.
                    {common_table_rules}
                    {common_signatures}
                    ==== AI STATUTE ANALYSIS ====
                    1. **Original Statute 121 Details:** - **Item Number Used:** [State the specific rule number you used].
                        - **Original Statute Text:** [Provide the EXACT quote/sentence directly from the attached Statute 121 PDF for this specific rule number].
                    2. **Justification:** [Explain exactly WHY this specific statute applies to the requested purchase. Relate the items being bought to the statute's wording].
                    3. **Similar Precedent from Sample Nondh:** [Find a similar past purchase in the uploaded 'Sample Nondh' context. List its Subject, Date, and the Statute Item Number it used to prove your choice is historically accurate].
                    4. **Rejected Alternative Statute:** [Find another statute item number from the PDF that looks similar but is INCORRECT (e.g., a rule for furniture instead of chemicals). Quote it and explicitly explain why it is NOT compatible with this purchase].
                    """

                    inputs = [sys_prompt, text_prompt]
                    if uploaded_image:
                        if uploaded_image.type == "application/pdf":
                            inputs.append({"mime_type": "application/pdf", "data": uploaded_image.getvalue()})
                        else:
                            inputs.append(Image.open(uploaded_image))
                        
                    response = model.generate_content(inputs)
                    res_text = response.text
                    
                    # Intercept and Split the AI Response safely
                    if "==== AI STATUTE ANALYSIS ====" in res_text:
                        parts = res_text.split("==== AI STATUTE ANALYSIS ====")
                        st.session_state['generated_nondh'] = parts[0].strip()
                        st.session_state['statute_analysis'] = parts[1].strip()
                    else:
                        st.session_state['generated_nondh'] = res_text.strip()
                        st.session_state['statute_analysis'] = ""
                    st.session_state['last_learning_context'] = learning
                    log_skill_runs(learning["skills"], "nondh_generation", text_prompt or "", "used")

                    st.success("સાદર નોંધ સફળતાપૂર્વક તૈયાર થઈ ગઈ છે!")
                    
                except Exception as e:
                    st.error(f"Error generating document: {e}")

    if 'generated_nondh' in st.session_state:
        
        # FEATURE ADDITION: Display Statute Analysis clearly separated from the Document Draft
        if 'statute_analysis' in st.session_state and st.session_state['statute_analysis']:
            with st.expander("🔍 Statute 121 Analysis & Justification (AI Reasoning)", expanded=True):
                st.info("આ વિભાગ ફક્ત તમારી જાણકારી માટે છે અને વર્ડ ડોક્યુમેન્ટ (DOCX) માં પ્રિન્ટ થશે નહીં.")
                st.markdown(st.session_state['statute_analysis'])

        learning_trace = st.session_state.get('last_learning_context', {})
        if learning_trace and (learning_trace.get("memories") or learning_trace.get("skills")):
            with st.expander("🧠 Learning Memory & Hermes Skills Used", expanded=False):
                for memory in learning_trace.get("memories", []):
                    st.markdown(f"**Memory #{memory['id']} ({memory['category']})**: {memory['title']}")
                    st.caption(compact_text(memory['content'], 500))
                for skill in learning_trace.get("skills", []):
                    st.markdown(f"**Skill #{skill['id']}**: {skill['name']} v{skill['version']}")
                    st.caption(compact_text(skill['goal'], 500))

        st.markdown("---")
        st.markdown("### ડ્રાફ્ટ એડિટિંગ (Smart Editor)")
        
        pre_text, df, post_text = parse_markdown_to_parts(st.session_state['generated_nondh'])
        
        edit_pre = st.text_area("ઉપરનું લખાણ:", pre_text, height=400)
        
        if not df.empty:
            st.markdown("#### સ્માર્ટ ટેબલ (Smart Table)")
            st.info("નોંધ: 'Required Quantity' અથવા 'Unit/Pkt Price' બદલશો તો 'Total Price' અને લખાણમાં રહેલ 'અંદાજિત ખર્ચ' આપોઆપ બદલાઈ જશે.")
            
            # ટેબલની કોલમની પહોળાઈ (width) સેટ કરવા માટેનું કન્ફિગરેશન
            custom_column_config = {
                "Sr. No.": st.column_config.TextColumn("Sr. No.", width="small"),
                "Details": st.column_config.TextColumn("Details", width="large"),
                "Required Quantity": st.column_config.TextColumn("Required Quantity", width="small"),
                "Available Pkt/Unit": st.column_config.TextColumn("Available Pkt/Unit", width="medium"),
                "Unit/Pkt Price": st.column_config.TextColumn("Unit/Pkt Price", width="medium"),
                "Total Price": st.column_config.TextColumn("Total Price", width="medium")
            }
            
            edited_df = st.data_editor(
                df, 
                num_rows="dynamic", 
                use_container_width=True,
                column_config=custom_column_config
            )
            
            # Intelligent Math Calculation using Regex to strip any accidental text
            if 'Required Quantity' in edited_df.columns and 'Unit/Pkt Price' in edited_df.columns and 'Total Price' in edited_df.columns:
                
                # Extract pure numbers safely
                req_qty = edited_df['Required Quantity'].astype(str).str.extract(r'(\d+\.?\d*)')[0].astype(float).fillna(0)
                unit_price = edited_df['Unit/Pkt Price'].astype(str).str.extract(r'(\d+\.?\d*)')[0].astype(float).fillna(0)
                
                # Perform the calculation
                edited_df['Total Price'] = (req_qty * unit_price).round(2)
                
                # Dynamic Grand Total Calculation
                grand_total_calc = edited_df['Total Price'].sum()
                st.success(f"**Grand Total (કુલ રકમ): ₹ {grand_total_calc:,.2f}**")
                
                # Automatically sync the paragraph text with the accurate Grand Total
                edit_pre = re.sub(r'(અંદાજિત ખર્ચ\s*).*?(\s*થનાર)', f'\g<1>{grand_total_calc:,.2f}\g<2>', edit_pre)
        else:
            edited_df = pd.DataFrame()
            st.info("આ નોંધમાં ટેબલની જરૂરિયાત જણાઈ નથી. (No table required for this note based on the context).")
            
        edit_post = st.text_area("નીચેનું લખાણ:", post_text, height=300)
        
        # Re-stitch using the custom markdown generator that handles the Grand Total
        final_document = f"{edit_pre}\n\n{df_to_markdown_with_total(edited_df)}\n{edit_post}" if not edited_df.empty else f"{edit_pre}\n\n{edit_post}"
        
        st.markdown("---")
        st.markdown("### દસ્તાવેજ પ્રીવ્યુ (Visual Preview - 20/80 Layout)")
        
        with st.container(border=True):
            prev_blank, prev_content = st.columns([2, 8]) 
            with prev_content:
                st.markdown(final_document)
        
        st.markdown("---")
        col_save, col_down = st.columns(2)
        with col_save:
            if st.button("આર્કાઇવમાં સેવ કરો (Save Nondh)"):
                subj = "No Subject"
                for line in final_document.split('\n'):
                    if "વિષય:" in line:
                        subj = line.replace("વિષય:", "").strip()
                        break
                
                # Save Nondh to DB and get the ID
                nondh_id = save_to_db(subj, final_document)
                st.session_state['recent_nondh_id'] = nondh_id
                
                # Automatically save the DOCX to the vault linked to this new Nondh
                docx_data = create_docx(final_document)
                save_file_to_vault(docx_data, f"Nondh_{nondh_id}_{datetime.date.today().strftime('%d_%m')}.docx", "Sadar Nondh Draft", nondh_id=nondh_id)
                with st.spinner("Creating learning suggestions from this completed Nondh..."):
                    suggestion_count = suggest_learning_from_nondh(
                        api_key,
                        st.session_state.get('generated_nondh', ''),
                        final_document,
                        subj,
                        nondh_id
                    )
                if suggestion_count:
                    st.info(f"{suggestion_count} learning suggestion(s) were added to Tab 6 for approval.")
                st.success("નોંધ અને ડ્રાફ્ટ સાચવી લેવામાં આવ્યા છે! (હવે તમે Tab 3 માં જઈ શકો છો)")
                
        with col_down:
            docx_data = create_docx(final_document)
            st.download_button(label="Download as Word (DOCX)",
                               data=docx_data,
                               file_name=f"Sadar_Nondh_{datetime.date.today().strftime('%d_%m_%Y')}.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        # --- NEW FEATURE: Delete Recent Nondh ---
        st.markdown("---")
        with st.expander("🗑️ તાજેતરની નોંધ રદ કરો (Delete Recent Nondh - Last 30 Days)", expanded=False):
            recent_nondhs = get_recent_nondhs(30)
            if not recent_nondhs:
                st.info("છેલ્લા 30 દિવસની કોઈ નોંધ ઉપલબ્ધ નથી.")
            else:
                for n_id, n_date, n_subj in recent_nondhs:
                    col_del1, col_del2 = st.columns([8, 2])
                    with col_del1:
                        st.write(f"**ID: {n_id}** | 🗓️ {n_date} | 📝 {n_subj}")
                    with col_del2:
                        confirm_delete_nondh = st.checkbox("Confirm", key=f"confirm_del_nondh_{n_id}")
                        if st.button("🗑️ Delete", key=f"del_nondh_{n_id}", type="secondary"):
                            if not confirm_delete_nondh:
                                st.warning("Please tick Confirm before deleting this Nondh.")
                            else:
                                delete_nondh(n_id)
                                st.error(f"નોંધ ID {n_id} સફળતાપૂર્વક રદ કરવામાં આવી છે!")
                                st.rerun()

with tab2:
    st.markdown("### 🗄️ જૂના રેકોર્ડ શોધો (Archive Search)")
    search_query = st.text_input("🔍 Smart Search (Type in English or Gujarati)")
    db_records = get_archives("All", "All")
    sample_records = search_sample_nondh("", "All", "All")
    all_records = db_records + sample_records

    def smart_search_gemini(query, records):
        if not query.strip(): return records
        if not api_key: return records
        learning = get_learning_context(query, ["nondh_style", "item_mapping", "general_workflow"], "archive_search")

        # Using a fast model for search
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-3.1-pro-preview')
        
        records_context = ""
        for i, record in enumerate(records):
            # FIXED: Handle 4-item database records properly
            if len(record) == 4:
                _, _, subject, content = record
            elif len(record) == 3: 
                date, subject, content = record
            elif len(record) == 2: 
                subject, content = record
            else: 
                continue 
            
            # INCREASED CONTEXT: Send up to 800 characters so AI can read the actual items
            clean_content = content[:800].replace('\n', ' ')
            records_context += f"ID: {i} | Subject: {subject} | Content: {clean_content}...\n"
            
        # ADVANCED PROMPT: Instruct the AI to act as a cross-lingual search engine
        prompt = f"""
        You are an intelligent bilingual search engine (English <-> Gujarati).
        The user is searching for: "{query}"
        
        Your Task: Find all record IDs that semantically match the user's query.
        - If the query is in English (e.g., "chemicals", "pins", "equipment"), translate the intent and find the matching Gujarati records.
        - If the query is in Gujarati, find the exact or contextually matching records.
        - Match based on meaning and translated keywords, not just exact text matches.
        
        Records Data:
        {records_context}

        Approved app memory/skills that may help search:
        {learning['prompt'] if learning['prompt'] else 'None'}

        Return ONLY a comma-separated list of matching IDs (e.g., 0,2,5).
        If absolutely no records match the meaning, return the exact word: NONE
        """

        try:
            log_skill_runs(learning["skills"], "archive_search", query, "used")
            response = model.generate_content(prompt)
            result = response.text.strip()
            
            if "NONE" in result.upper() or not result: 
                return []
                
            # Safely extract only numbers using regex (prevents crashes if AI adds extra words)
            import re
            matched_ids_str = re.findall(r'\d+', result)
            matched_ids = [int(x) for x in matched_ids_str]
            
            # Return only valid records
            return [records[i] for i in matched_ids if i < len(records)]
            
        except Exception as e:
            return records

    display_records = smart_search_gemini(search_query, all_records) if search_query else all_records
    if display_records:
        st.success(f"કુલ {len(display_records)} રેકોર્ડ મળ્યા.")
        for idx, record in enumerate(display_records):
            if len(record) == 4:
                nondh_id, date, subject, content = record[0], record[1], record[2], record[3]
            else:
                continue # Skip invalid records
            
            with st.expander(f"ID #{nondh_id if nondh_id else 'Ref'} | {date} - {subject}"):
                st.markdown(content)
                st.download_button("Download Nondh (Word)", data=create_docx(content), file_name=f"Nondh_{nondh_id}.docx", key=f"dl_{idx}") 
                
                # --- Unified View of Vault Files linked to this Nondh ---
                if nondh_id:
                    vault_files = get_vault_files_by_nondh(nondh_id)
                    if vault_files:
                        st.markdown("---")
                        st.markdown("#### 📂 જોડાયેલ દસ્તાવેજો (Linked Vault Documents)")
                        for f_name, f_path, u_date, d_type, desc in vault_files:
                            col1, col2 = st.columns([8, 2])
                            with col1: st.caption(f"**{d_type}**: {f_name} ({u_date}) - {desc}")
                            with col2:
                                file_data_t2 = load_vault_file_bytes(f_path)
                                if file_data_t2:
                                        st.download_button("⬇️", data=file_data_t2, file_name=f_name, key=f"dl_v_{f_path}")
                                else: 
                                    st.error("Missing")

with tab3:
    st.markdown("### 📝 ખરીદી હુકમ બનાવો (Generate Purchase Order)")
    st.info("નોંધ મંજૂર થયા પછી સપ્લાયરને ખરીદીનો ઓર્ડર મોકલવા માટે અહીં વિગતો ભરો.")
    
    st.markdown("#### ૧. મંજૂર થયેલ નોંધ પસંદ કરો (Select Approved Nondh)")
    db_records = get_archives("All", "All")
    
    # --- NEW: Check the database for Nondh IDs that already have a PO ---
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("SELECT nondh_id FROM purchase_orders WHERE nondh_id IS NOT NULL")
    used_nondh_ids = [row[0] for row in c.fetchall()]
    conn.close()

    options = ["-- જાતે માહિતી ભરો (Manual Entry) --"]
    record_dict = {}
    for idx, row in enumerate(db_records):
        if len(row) == 4:
            nondh_id = row[0]
            # ONLY add to dropdown if it hasn't been used yet!
            if nondh_id not in used_nondh_ids:
                label = f"[{nondh_id}] {row[1]} - {row[2]}" 
                record_dict[label] = row[3] 
                options.append(label)
                
    selected_nondh = st.selectbox("અગાઉ સેવ કરેલ નોંધ પસંદ કરો:", options)
    
    # --- Extract Nondh ID early so we can link the upload to it ---
    current_nondh_id = None
    if selected_nondh != "-- જાતે માહિતી ભરો (Manual Entry) --":
        import re
        match = re.search(r'\[(\d+)\]', selected_nondh)
        if match:
            current_nondh_id = int(match.group(1))

    st.markdown("#### ૨. સહી કરેલ નોંધ અપલોડ કરો (Upload Signed Nondh - Optional)")
    uploaded_nondh = st.file_uploader("મંજૂર થયેલ/સહીવાળી નોંધ અપલોડ કરો:", type=["pdf", "jpg", "jpeg", "png"], key="nondh_up")
    
    if uploaded_nondh:
        if not current_nondh_id:
            st.warning("ફાઈલ અપલોડ કરતા પહેલા ઉપરથી 'અગાઉ સેવ કરેલ નોંધ' પસંદ કરો!")
        else:
            # Check if already uploaded this session to prevent duplicates!
            if st.session_state.get(f"uploaded_nondh_{current_nondh_id}") != uploaded_nondh.name:
                file_bytes = uploaded_nondh.getbuffer()
                # Use explicit nondh_id keyword to link it perfectly!
                save_file_to_vault(file_bytes, uploaded_nondh.name, "Signed Nondh", nondh_id=current_nondh_id, description="Auto-uploaded from Tab 3")
                st.session_state[f"uploaded_nondh_{current_nondh_id}"] = uploaded_nondh.name
                st.success("સહી કરેલ નોંધ વોલ્ટમાં સાચવી લેવાઈ છે!")

    st.markdown("---")
    st.markdown("#### ૩. સપ્લાયર અને ઓર્ડરની વિગત (Supplier & Order Details)")
    col_po1, col_po2 = st.columns(2)
    with col_po1:
        vendor_name = st.text_input("સપ્લાયરનું નામ (Vendor Name)", value="DUTT ENTERPRISE")
        outward_no = st.text_input("જાવક નંબર (Outward No.)", value="139")
    with col_po2:
        vendor_address = st.text_area("સપ્લાયરનું સરનામું", value="A/5, Krishna complex, borsad chokadi,\nAnand sojitra road, Anand 388 001", height=110)
        po_date = st.date_input("તારીખ (Date)", value=datetime.date.today())
        
    default_df = pd.DataFrame(columns=["Details", "Required Quantity", "Available Pkt/Unit", "Unit/Pkt Price", "Total Price"])
    if selected_nondh != "-- જાતે માહિતી ભરો (Manual Entry) --":
        _, session_df, _ = parse_markdown_to_parts(record_dict[selected_nondh])
        if not session_df.empty: default_df = session_df
    
    po_df = st.data_editor(default_df, num_rows="dynamic", use_container_width=True, key="po_editor")
    
   # ડાઉનલોડ બટન અને ડેટાબેઝ સેવ કરવાનું નવું લોજીક
    # ડાઉનલોડ બટન અને ડેટાબેઝ સેવ કરવાનું નવું લોજીક
    if vendor_name and not po_df.empty:
        formatted_date = po_date.strftime("%d.%m.%Y")
        grand_total = coerce_amount(pd.to_numeric(po_df['Total Price'], errors='coerce').fillna(0).sum())
        
        current_nondh_id = None
        if selected_nondh != "-- જાતે માહિતી ભરો (Manual Entry) --":
            import re
            match = re.search(r'\[(\d+)\]', selected_nondh)
            if match:
                current_nondh_id = int(match.group(1))
        
        # 1. પહેલાં વર્ડ ફાઈલ બેકગ્રાઉન્ડમાં તૈયાર કરી લો
# 1. પહેલાં વર્ડ ફાઈલ બેકગ્રાઉન્ડમાં તૈયાર કરી લો
        po_docx = create_purchase_order_docx(vendor_name, vendor_address, outward_no, formatted_date, po_df)
        
        st.markdown("---")
        st.info("સૂચના: પહેલા 'સ્ટેપ ૧' બટન પર ક્લિક કરીને ફાઈલ ડાઉનલોડ કરો, ત્યારબાદ જ 'સ્ટેપ ૨' પર ક્લિક કરો.")
        
        col_btn1, col_btn2 = st.columns(2)
        
        with col_btn1:
            st.download_button(
                label="૧. 📄 ખરીદી હુકમ ડાઉનલોડ કરો (Step 1: Download PO)", 
                data=po_docx, 
                file_name=f"PO_{vendor_name}.docx"
            )

        with col_btn2:
            if st.button("૨. ➡️ પેમેન્ટ માટે આગળ મોકલો (Step 2: Send to Tab 4)", type="primary"):
                if not vendor_name.strip():
                    st.error("Please enter the vendor name before sending this PO to payment.")
                elif not str(outward_no).strip():
                    st.error("Please enter the outward number before sending this PO to payment.")
                elif grand_total <= 0:
                    st.error("PO total must be greater than zero before sending this PO to payment.")
                elif current_nondh_id and get_po_for_nondh(current_nondh_id):
                    st.warning("A PO already exists for this Nondh. It was not added again.")
                else:
                    save_po_to_db(current_nondh_id, vendor_name, outward_no, formatted_date, grand_total)
                    st.success("ઓર્ડર સફળતાપૂર્વક Tab 4 (Bill Payment) માં મોકલી દેવાયો છે!")
                    st.rerun()
# --- TAB 4 (Bill Payment ONLY) ---
with tab4:
    st.markdown("### 💳 બિલ પેમેન્ટ ફોર્મ (Bill Payment Form)")
    st.info("જે ખરીદીના હુકમ (Purchase Orders) માટે બિલ ચૂકવવાનું બાકી છે, તે જ અહીં દેખાશે.")
    
    unfinished_pos = get_unfinished_pos(('Unfinished',))
    
    if not unfinished_pos:
        st.success("હાલમાં કોઈ બિલ પેમેન્ટ બાકી નથી! (No unfinished purchase orders).")
    else:
        po_dict_tab4 = {}
        po_options_tab4 = []
        for po in unfinished_pos:
            # FIXED: Added nondh_id_t4 to catch all 6 items from the database
            po_id, nondh_id_t4, v_name, o_no, p_date, amt = po
            amt = coerce_amount(amt)
            label = f"PO #{o_no} - {v_name} - ₹{format_amount(amt)} ({p_date})"
            po_options_tab4.append(label)
            po_dict_tab4[label] = (po_id, nondh_id_t4, v_name, o_no, p_date, amt)
            
        selected_po_label_t4 = st.selectbox("પેમેન્ટ ફોર્મ માટે ઓર્ડર પસંદ કરો (Select Pending PO):", po_options_tab4, key="po_tab4")
        
        if selected_po_label_t4:
            # FIXED: Added nondh_id_t4 here as well
            po_id, nondh_id_t4, v_name, o_no, p_date, amt = po_dict_tab4[selected_po_label_t4]
            
            if (
                st.session_state.get("current_po_id_t4") != po_id
                or "amt_t4" not in st.session_state
                or "ext_amt" not in st.session_state
                or "ext_bill_no" not in st.session_state
                or "ext_words" not in st.session_state
                or "last_invoice" not in st.session_state
            ):
                st.session_state.current_po_id_t4 = po_id
                st.session_state.ext_bill_no = "INV-"
                st.session_state.ext_amt = coerce_amount(amt)
                st.session_state.amt_t4 = coerce_amount(amt)
                st.session_state.ext_words = ""
                st.session_state.last_invoice = None

            st.markdown("#### ઇન્વોઇસ અને બજેટની વિગતો (Invoice Details)")
            
            col_b1, col_b2 = st.columns(2)
            with col_b1:
                budget_head = st.text_input("Budget Head No.", value="303/2092 (AINP on Agril Acarology)", key="bh_t4")
                bill_no = st.text_input("ઇન્વોઇસ/બિલ નંબર (Vendor Bill No.)", value=st.session_state.ext_bill_no)
                invoice_upload = st.file_uploader("પાર્ટીનું બિલ અપલોડ કરો (Upload Vendor Invoice PDF/Img)", type=["pdf", "jpg", "png"])
                
                if invoice_upload:
                    # Prevent duplicate uploads on app rerun
                    if st.session_state.get(f"uploaded_inv_{po_id}") != invoice_upload.name:
                        file_bytes = invoice_upload.getbuffer()
                        
                        # Use explicit nondh_id keyword to link it perfectly!
                        save_file_to_vault(file_bytes, invoice_upload.name, "Party Invoice", nondh_id=nondh_id_t4, description=f"Auto-uploaded from Tab 4 for PO #{o_no}")
                        st.session_state[f"uploaded_inv_{po_id}"] = invoice_upload.name
                        st.success("ઇન્વોઇસ વોલ્ટમાં Nondh ID સાથે સેવ થઈ ગયું!")

                    if invoice_upload.name != st.session_state.last_invoice and api_key:
                        with st.spinner("AI દ્વારા બિલની વિગતો વાંચવામાં આવી રહી છે... (Extracting...)"):
                            try:
                                import json
                                import re  # ટેક્સ્ટ સાફ કરવા માટે
                                genai.configure(api_key=api_key)
                                model = genai.GenerativeModel('gemini-3.1-pro-preview')
                                learning = get_learning_context(
                                    f"{v_name} {o_no} {amt} {invoice_upload.name}",
                                    ["invoice_extraction", "bill_payment", "vendor_default", "general_workflow"],
                                    "invoice_extraction"
                                )

                                prompt = f"""
                                You are an intelligent accounting AI. The approved Purchase Order (PO) amount for this transaction is ₹{amt}.
                                Approved memories and Hermes-style skills:
                                {learning['prompt'] if learning['prompt'] else 'None'}

                                Carefully analyze the uploaded invoice and extract the following:
                                
                                1. Invoice/Bill Number.
                                2. Final Payable Amount: Look for terms like 'Grand Total', 'Invoice Total', 'Net Payable', or 'Total Amount'. Use your intelligence to understand the invoice structure and identify the final amount including taxes. Logically compare it with the PO amount (₹{amt}) to ensure you pick the correct total. Return it as a PURE NUMBER WITHOUT COMMAS (e.g. 3894.00).
                                3. Final Payable Amount in English words.
                                
                                Return ONLY a valid JSON object in this exact format:
                                {{"bill_no": "INV-123", "amount": 1234.50, "amount_words": "One Thousand Two Hundred..."}}
                                """
                                
                                # --- મોટો સુધારો: PyPDF2 કાઢીને PDF સીધી જ Gemini ને આપો ---
                                if invoice_upload.type == "application/pdf":
                                    doc_parts = [{"mime_type": "application/pdf", "data": invoice_upload.getvalue()}]
                                    log_skill_runs(learning["skills"], "invoice_extraction", invoice_upload.name, "used")
                                    response = model.generate_content([prompt, doc_parts[0]])
                                else:
                                    img = Image.open(invoice_upload)
                                    log_skill_runs(learning["skills"], "invoice_extraction", invoice_upload.name, "used")
                                    response = model.generate_content([prompt, img])
                                
                                res_text = response.text.strip().replace("```json", "").replace("```", "")
                                data = json.loads(res_text)
                                
                                # --- બીજો સુધારો: રકમમાંથી કોમા (,) અને અન્ય ચિન્હો સાફ કરો ---
                                raw_amount = coerce_amount(data.get("amount", amt), amt)
                                
                                st.session_state.ext_bill_no = str(data.get("bill_no", "INV-"))
                                st.session_state.ext_amt = raw_amount
                                st.session_state.amt_t4 = raw_amount
                                st.session_state.ext_words = str(data.get("amount_words", ""))
                                st.session_state.last_invoice = invoice_upload.name
                                st.rerun() 
                            except Exception as e:
                                st.warning(f"આપમેળે વિગત મેળવવામાં ભૂલ: {e}. કૃપા કરીને જાતે ભરો.")

            with col_b2:
                bill_date = st.date_input("ઇન્વોઇસની તારીખ (Bill Date)", value=datetime.date.today())
                final_amt = st.number_input("ચૂકવવા પાત્ર રકમ (Amount to Pay)", min_value=0.0, step=1.0, format="%.2f", key="amt_t4")
                
                # --- NEW FEATURE: Auto-fill English words based on Amount ---
                col_eng1, col_eng2 = st.columns([3, 1])
                with col_eng1:
                    amount_words = st.text_input("રકમ શબ્દોમાં (Amount in Words - English)", value=st.session_state.ext_words, placeholder="e.g., Four Thousand Two Hundred Forty Eight")
                with col_eng2:
                    st.write("") 
                    if st.button("✨ AI થી ભરો", key="auto_fill_eng_t4"):
                        if api_key:
                            with st.spinner("Converting to words..."):
                                try:
                                    genai.configure(api_key=api_key)
                                    model = genai.GenerativeModel('gemini-3.1-pro-preview')
                                    learning = get_learning_context(str(final_amt), ["bill_payment", "general_workflow"], "amount_words_english")
                                    prompt = f"""Convert the number {final_amt} into English words (capitalize the first letter of each word).
Approved memories and Hermes-style skills:
{learning['prompt'] if learning['prompt'] else 'None'}
Return ONLY the text, nothing else. Example: for 3956.50 return 'Three Thousand Nine Hundred Fifty Six And Fifty Paise'."""
                                    log_skill_runs(learning["skills"], "amount_words_english", str(final_amt), "used")
                                    res = model.generate_content(prompt)
                                    st.session_state.ext_words = res.text.strip()
                                    st.rerun()
                                except Exception as e:
                                    st.error("AI Error.")
                        else:
                            st.warning("API Key is required!")
            
            st.markdown("---")
            if st.button("📄 Generate Bill Payment Form"):
                if not amount_words: st.error("Please enter the amount in words!")
                else:
                    bp_docx = create_bill_payment_form(budget_head, bill_no, bill_date.strftime("%d/%m/%Y"), v_name, final_amt, amount_words)
                    st.download_button(
                        label="Download Bill Payment Form", 
                        data=bp_docx, 
                        file_name=f"Payment_Form_{v_name}.docx",
                        on_click=mark_po_as_payment_generated,
                        args=(po_id,)
                    )

# --- TAB 5 (Bill Pasting & Mark Paid ONLY) ---
with tab5:
    st.markdown("### 📑 બિલ પેસ્ટિંગ અને પ્રમાણપત્ર (Bill Pasting Form)")
    
    if "auto_guj_words" not in st.session_state:
        st.session_state.auto_guj_words = ""

    unfinished_pos_t5 = get_unfinished_pos(('Unfinished', 'Payment_Generated'))
    
    if not unfinished_pos_t5:
        st.success("હાલમાં કોઈ બિલ પેમેન્ટ બાકી નથી! (No unfinished purchase orders).")
    else:
        po_dict_tab5 = {}
        po_options_tab5 = []
        for po in unfinished_pos_t5:
            # FIXED: Added nondh_id_t5 to catch all 6 items from the database
            po_id, nondh_id_t5, v_name, o_no, p_date, amt = po
            amt = coerce_amount(amt)
            label = f"PO #{o_no} - {v_name} - ₹{format_amount(amt)} ({p_date})"
            po_options_tab5.append(label)
            po_dict_tab5[label] = (po_id, nondh_id_t5, v_name, o_no, p_date, amt)
            
        selected_po_label_t5 = st.selectbox("પેસ્ટિંગ ફોર્મ માટે ઓર્ડર પસંદ કરો:", po_options_tab5, key="po_tab5")
        
        if selected_po_label_t5:
            # FIXED: Added nondh_id_t5 here as well
            po_id_t5, nondh_id_t5, v_name_t5, o_no_t5, p_date_t5, amt_t5 = po_dict_tab5[selected_po_label_t5]

            col_p1, col_p2 = st.columns(2)
            with col_p1:
                budget_head_pst = st.text_input("Budget Head No.", value="303/2092 (AINP on Agril Acarology)", key="bh_t5")
                grant_year = st.text_input("ફાળવેલ ગ્રાન્ટ વર્ષ (Grant Year)", value="", placeholder="હાથેથી લખવા માટે ખાલી છોડી દો")
                party_name_pst = st.text_input("પાર્ટીનું નામ (Party Name)", value=v_name_t5, key="party_t5")
            with col_p2:
                # --- નવો સુધારો: Streamlit ની મેમરીને સીધી અપડેટ કરવા માટેનું લોજીક ---
                # જો Tab 4 માં આ જ ઓર્ડર ખૂલ્યો હોય અને ત્યાં રકમ સેટ હોય, તો Tab 5 ની મેમરી ફરજિયાત ઓવરરાઈટ કરો
                amount_from_tab4 = st.session_state.get("current_po_id_t4") == po_id_t5 and "amt_t4" in st.session_state
                amount_for_t5 = coerce_amount(st.session_state.amt_t4) if amount_from_tab4 else coerce_amount(amt_t5)
                if st.session_state.get("current_po_id_t5") != po_id_t5:
                    st.session_state.current_po_id_t5 = po_id_t5
                    st.session_state['amt_t5'] = amount_for_t5
                elif amount_from_tab4:
                    st.session_state['amt_t5'] = amount_for_t5
                
                # નોંધ: અહીથી `value=...` કાઢી નાખ્યું છે, કારણ કે તે સીધું `key="amt_t5"` ની મેમરીમાંથી જ લેટેસ્ટ રકમ ખેંચી લેશે
                final_amt_pst = st.number_input("બીલની કુલ રકમ (Amount)", min_value=0.0, step=1.0, format="%.2f", key="amt_t5")
                
                # --- બટન વગર આપોઆપ ગુજરાતી અનુવાદ (Auto-Translate) ---
                @st.cache_data(show_spinner=False)
                def get_gujarati_words_auto(amount, key, learning_prompt=""):
                    if not key or amount == 0: return ""
                    try:
                        genai.configure(api_key=key)
                        model = genai.GenerativeModel('gemini-3.1-flash-lite-preview')
                        prompt = f"""Translate the number {amount} into Gujarati words.
Approved memories and Hermes-style skills:
{learning_prompt if learning_prompt else 'None'}
Return ONLY the Gujarati translation. Example: for 3956 return 'ત્રણ હજાર નવસો છપ્પન'."""
                        res = model.generate_content(prompt)
                        return res.text.strip()
                    except:
                        return ""

                learning_t5 = get_learning_context(
                    f"{party_name_pst} {final_amt_pst}",
                    ["bill_pasting", "register_rule", "bill_payment", "general_workflow"],
                    "amount_words_gujarati"
                )
                skill_log_key = f"guj_skill_logged_{po_id_t5}_{final_amt_pst}"
                if not st.session_state.get(skill_log_key):
                    log_skill_runs(learning_t5["skills"], "amount_words_gujarati", str(final_amt_pst), "used")
                    st.session_state[skill_log_key] = True
                auto_gujarati_text = get_gujarati_words_auto(final_amt_pst, api_key, learning_t5["prompt"])
                amt_words_guj = st.text_input("રકમ શબ્દોમાં (ગુજરાતીમાં)", value=auto_gujarati_text, placeholder="દા.ત., ત્રણ હજાર નવસો છપ્પન")
                            
            st.markdown("#### 📝 મંજુરીની વિગતો (Approval Details - મુદ્દા નં. ૧)")
            col_a1, col_a2, col_a3 = st.columns(3)
            with col_a1:
                item_no_pst = st.text_input("સ્ટેચ્યુટ નં. ૧૨૧ ની આઇટમ નં", value="", placeholder="દા.ત. 14")
            with col_a2:
                approval_no_pst = st.text_input("મંજુરી નં. (Approval No.)", value="", placeholder="દા.ત. ACN/123/2026")
            with col_a3:
                approval_date_pst = st.text_input("મંજુરી તારીખ (Approval Date)", value="", placeholder="DD/MM/YYYY")
                
            st.markdown("#### 📝 રજીસ્ટર અને નોંધની વિગતો (Register Details - મુદ્દા નં. ૩ અને ૮)")
            col_r1, col_r2 = st.columns(2)
            with col_r1:
                reg_type = st.selectbox("કયા રજીસ્ટરમાં નોંધ કરી? (મુદ્દા નં. ૩)", ["ચીજવસ્તુ વપરાશ (કન્ઝયુમેબલ)", "ડેડસ્ટોક", "સ્ટોર રોજમેળ", "ટેલીફોન", "સ્ટેમ્પ", "સ્ટેશનરી", "પરચુરણ માલ સામાન", "રીપેરીંગ"])
                reg_page_no = st.text_input("રજીસ્ટર પાના નં. (Register Page No.)", value="")
            with col_r2:
                bill_reg_date = st.date_input("બીલ રજીસ્ટરમાં નોંધની તારીખ (મુદ્દા નં. ૮)", value=datetime.date.today())
                bill_reg_page_no = st.text_input("બીલ રજી પાના નં. (Bill Reg. Page No.)", value="")
                bill_reg_sr_no = st.text_input("બીલ રજી અનુ. નં. (Bill Reg. Serial No.)", value="")
                
            st.markdown("---")
            col_btn_pst1, col_btn_pst2 = st.columns(2)
            
            with col_btn_pst1:
                if st.button("📑 Generate Exact Bill Pasting Form"):
                    if not amt_words_guj:
                        st.error("કૃપા કરીને રકમ શબ્દોમાં (ગુજરાતીમાં) લખો અથવા 'AI થી ભરો' બટન દબાવો.")
                    elif not reg_page_no or not bill_reg_page_no or not bill_reg_sr_no:
                        st.warning("કૃપા કરીને રજીસ્ટરના પાના નંબર અને અનુક્રમ નંબર ભરો.")
                    else:
                        bill_reg_date_str = bill_reg_date.strftime("%d/%m/%Y")
                        
                        # Added new parameters to function call
                        pst_docx = create_bill_pasting_form(
                            budget_head_pst, grant_year, party_name_pst, final_amt_pst, 
                            amt_words_guj, reg_type, reg_page_no, bill_reg_date_str, 
                            bill_reg_page_no, bill_reg_sr_no, item_no_pst, approval_no_pst, approval_date_pst
                        )
                        st.download_button("Download Pasting Form", data=pst_docx, file_name=f"Pasting_Form_{v_name_t5}.docx")
            
            with col_btn_pst2:
                if st.button("✅ બિલ પેમેન્ટ પૂરું કરો (Mark as Paid)", key="mark_paid"):
                    mark_po_as_paid(po_id_t5)
                    st.success("ઓર્ડર પેમેન્ટ લિસ્ટમાંથી દૂર કરવામાં આવ્યો છે! રિફ્રેશ કરો.")
                    st.rerun()
# --- TAB 6 (Digital Vault / Archive) ---
# --- TAB 6 (DIGITAL VAULT & PAYMENT CLOSURE) ---
with tab6:
    st.markdown("### 🗄️ ડિજિટલ વોલ્ટ અને પેમેન્ટ ક્લોઝર (Vault & Payment Closure)")

    with st.expander("🧠 Learning Memory & Hermes-Style Skills", expanded=False):
        mem_tab, mem_sug_tab, skill_tab, skill_sug_tab, run_tab = st.tabs([
            "Approved Memories", "Memory Suggestions", "Approved Skills", "Skill Suggestions", "Skill Runs"
        ])

        with mem_tab:
            st.markdown("#### Add Memory")
            with st.form("add_memory_form"):
                m_col1, m_col2, m_col3 = st.columns([2, 3, 1])
                with m_col1:
                    new_mem_category = st.selectbox("Category", LEARNING_CATEGORIES, key="new_mem_category")
                    new_mem_priority = st.number_input("Priority", min_value=1, max_value=10, value=5, step=1, key="new_mem_priority")
                with m_col2:
                    new_mem_title = st.text_input("Title", key="new_mem_title")
                    new_mem_keywords = st.text_input("Keywords", key="new_mem_keywords")
                with m_col3:
                    new_mem_active = st.checkbox("Active", value=True, key="new_mem_active")
                new_mem_content = st.text_area("Memory Content", key="new_mem_content", height=120)
                if st.form_submit_button("Save Memory"):
                    if not new_mem_title.strip() or not new_mem_content.strip():
                        st.error("Memory title and content are required.")
                    else:
                        save_memory(new_mem_category, new_mem_title, new_mem_content, new_mem_keywords, new_mem_priority, int(new_mem_active), "manual", "")
                        st.success("Memory saved.")
                        st.rerun()

            st.markdown("#### Existing Memories")
            memories = list_memories(include_inactive=True)
            if not memories:
                st.info("No memories saved yet.")
            for memory in memories:
                mem_id, category, title, content, keywords, priority, active, source_type, source_id, created_at, updated_at = memory
                status = "Active" if active else "Inactive"
                with st.expander(f"#{mem_id} [{status}] {title}", expanded=False):
                    e_col1, e_col2 = st.columns([1, 2])
                    with e_col1:
                        edit_category = st.selectbox("Category", LEARNING_CATEGORIES, index=LEARNING_CATEGORIES.index(category) if category in LEARNING_CATEGORIES else 0, key=f"mem_cat_{mem_id}")
                        edit_priority = st.number_input("Priority", min_value=1, max_value=10, value=int(priority or 5), step=1, key=f"mem_pri_{mem_id}")
                        edit_active = st.checkbox("Active", value=bool(active), key=f"mem_act_{mem_id}")
                    with e_col2:
                        edit_title = st.text_input("Title", value=title or "", key=f"mem_title_{mem_id}")
                        edit_keywords = st.text_input("Keywords", value=keywords or "", key=f"mem_kw_{mem_id}")
                    edit_content = st.text_area("Content", value=content or "", height=120, key=f"mem_content_{mem_id}")
                    st.caption(f"Source: {source_type or '-'} {source_id or ''} | Updated: {updated_at or '-'}")
                    b1, b2, b3 = st.columns(3)
                    with b1:
                        if st.button("Update", key=f"mem_update_{mem_id}"):
                            update_memory(mem_id, edit_category, edit_title, edit_content, edit_keywords, edit_priority, int(edit_active))
                            st.success("Memory updated.")
                            st.rerun()
                    with b2:
                        if st.button("Deactivate" if active else "Activate", key=f"mem_toggle_{mem_id}"):
                            set_memory_active(mem_id, 0 if active else 1)
                            st.rerun()
                    with b3:
                        if st.button("Delete", key=f"mem_delete_{mem_id}"):
                            delete_memory(mem_id)
                            st.warning("Memory deleted.")
                            st.rerun()

        with mem_sug_tab:
            suggestions = list_memory_suggestions("Pending")
            if not suggestions:
                st.info("No pending memory suggestions.")
            for suggestion in suggestions:
                sug_id, category, title, content, keywords, priority, reason, source_type, source_id, source_snapshot, status, created_at, updated_at = suggestion
                with st.expander(f"Suggestion #{sug_id}: {title}", expanded=False):
                    st.caption(f"Category: {category} | Priority: {priority} | Source: {source_type} {source_id}")
                    st.write(reason or "No reason provided.")
                    st.markdown(content)
                    st.caption(f"Keywords: {keywords}")
                    with st.expander("Source Snapshot", expanded=False):
                        st.text(source_snapshot or "")
                    a_col, r_col = st.columns(2)
                    with a_col:
                        if st.button("Approve Memory", key=f"approve_mem_sug_{sug_id}"):
                            approve_memory_suggestion(sug_id)
                            st.success("Memory approved.")
                            st.rerun()
                    with r_col:
                        if st.button("Reject", key=f"reject_mem_sug_{sug_id}"):
                            reject_memory_suggestion(sug_id)
                            st.warning("Memory suggestion rejected.")
                            st.rerun()

        with skill_tab:
            st.markdown("#### Add Hermes-Style Skill")
            with st.form("add_skill_form"):
                s_col1, s_col2 = st.columns([2, 1])
                with s_col1:
                    new_skill_name = st.text_input("Skill Name", key="new_skill_name")
                    new_skill_keywords = st.text_input("Trigger Keywords", key="new_skill_keywords")
                with s_col2:
                    new_skill_priority = st.number_input("Priority", min_value=1, max_value=10, value=5, step=1, key="new_skill_priority")
                    new_skill_active = st.checkbox("Active", value=True, key="new_skill_active")
                new_skill_goal = st.text_area("Goal", key="new_skill_goal", height=80)
                new_skill_steps = st.text_area("Steps", key="new_skill_steps", height=120)
                new_skill_examples = st.text_area("Examples", key="new_skill_examples", height=80)
                new_skill_validation = st.text_area("Validation Rules", key="new_skill_validation", height=80)
                if st.form_submit_button("Save Skill"):
                    if not new_skill_name.strip() or not new_skill_goal.strip() or not new_skill_steps.strip():
                        st.error("Skill name, goal, and steps are required.")
                    else:
                        save_skill(new_skill_name, new_skill_keywords, new_skill_goal, new_skill_steps, new_skill_examples, new_skill_validation, new_skill_priority, int(new_skill_active), 1)
                        st.success("Skill saved.")
                        st.rerun()

            st.markdown("#### Existing Skills")
            skills = list_skills(include_inactive=True)
            if not skills:
                st.info("No skills saved yet.")
            for skill in skills:
                skill_id, name, trigger_keywords, goal, steps, examples, validation_rules, priority, active, version, created_at, updated_at = skill
                status = "Active" if active else "Inactive"
                with st.expander(f"#{skill_id} [{status}] {name} v{version}", expanded=False):
                    sk_col1, sk_col2 = st.columns([2, 1])
                    with sk_col1:
                        edit_name = st.text_input("Name", value=name or "", key=f"skill_name_{skill_id}")
                        edit_keywords = st.text_input("Trigger Keywords", value=trigger_keywords or "", key=f"skill_kw_{skill_id}")
                    with sk_col2:
                        edit_priority = st.number_input("Priority", min_value=1, max_value=10, value=int(priority or 5), step=1, key=f"skill_pri_{skill_id}")
                        edit_version = st.number_input("Version", min_value=1, value=int(version or 1), step=1, key=f"skill_ver_{skill_id}")
                        edit_active = st.checkbox("Active", value=bool(active), key=f"skill_act_{skill_id}")
                    edit_goal = st.text_area("Goal", value=goal or "", height=80, key=f"skill_goal_{skill_id}")
                    edit_steps = st.text_area("Steps", value=steps or "", height=120, key=f"skill_steps_{skill_id}")
                    edit_examples = st.text_area("Examples", value=examples or "", height=80, key=f"skill_examples_{skill_id}")
                    edit_validation = st.text_area("Validation Rules", value=validation_rules or "", height=80, key=f"skill_validation_{skill_id}")
                    st.caption(f"Updated: {updated_at or '-'}")
                    b1, b2, b3 = st.columns(3)
                    with b1:
                        if st.button("Update Skill", key=f"skill_update_{skill_id}"):
                            update_skill(skill_id, edit_name, edit_keywords, edit_goal, edit_steps, edit_examples, edit_validation, edit_priority, int(edit_active), edit_version)
                            st.success("Skill updated.")
                            st.rerun()
                    with b2:
                        if st.button("Deactivate" if active else "Activate", key=f"skill_toggle_{skill_id}"):
                            set_skill_active(skill_id, 0 if active else 1)
                            st.rerun()
                    with b3:
                        if st.button("Delete Skill", key=f"skill_delete_{skill_id}"):
                            delete_skill(skill_id)
                            st.warning("Skill deleted.")
                            st.rerun()

        with skill_sug_tab:
            suggestions = list_skill_suggestions("Pending")
            if not suggestions:
                st.info("No pending skill suggestions.")
            for suggestion in suggestions:
                sug_id, name, trigger_keywords, goal, steps, examples, validation_rules, priority, reason, source_type, source_id, source_snapshot, status, created_at, updated_at = suggestion
                with st.expander(f"Skill Suggestion #{sug_id}: {name}", expanded=False):
                    st.caption(f"Priority: {priority} | Source: {source_type} {source_id}")
                    st.write(reason or "No reason provided.")
                    st.markdown(f"**Goal:** {goal}")
                    st.markdown(f"**Steps:** {steps}")
                    st.markdown(f"**Validation:** {validation_rules}")
                    st.caption(f"Triggers: {trigger_keywords}")
                    with st.expander("Examples / Source", expanded=False):
                        st.text(f"{examples or ''}\n\n{source_snapshot or ''}")
                    a_col, r_col = st.columns(2)
                    with a_col:
                        if st.button("Approve Skill", key=f"approve_skill_sug_{sug_id}"):
                            approve_skill_suggestion(sug_id)
                            st.success("Skill approved.")
                            st.rerun()
                    with r_col:
                        if st.button("Reject", key=f"reject_skill_sug_{sug_id}"):
                            reject_skill_suggestion(sug_id)
                            st.warning("Skill suggestion rejected.")
                            st.rerun()

        with run_tab:
            runs = list_skill_runs(75)
            if not runs:
                st.info("No skill run history yet.")
            else:
                for run_id, skill_name, workflow, context_summary, outcome, source_id, created_at in runs:
                    st.markdown(f"**#{run_id} {skill_name}** | {workflow} | {outcome}")
                    st.caption(f"{created_at} | Source: {source_id or '-'} | {context_summary}")

    # Section A: Mark as Paid
    with st.expander("✅ બાકી પેમેન્ટ ક્લિયર કરો (Pending Payments to Mark as Paid)", expanded=False):
        pending_pos = get_unfinished_pos(('Unfinished', 'Payment_Generated'))
        if not pending_pos: st.info("કોઈ પેમેન્ટ બાકી નથી.")
        else:
            p_dict = {f"PO #{p[3]} - {p[2]} (₹{format_amount(p[5])})": p for p in pending_pos}
            sel_pay = st.selectbox("પેમેન્ટ થયેલ ઓર્ડર પસંદ કરો:", list(p_dict.keys()))
            if sel_pay:
                po_data = p_dict[sel_pay]
                col_pay1, col_pay2, col_pay3 = st.columns(3)
                with col_pay1: pay_info = st.text_input("પેમેન્ટની વિગત (UTR / Cheque No. / Date) - Optional")
                with col_pay2:
                    st.write("")
                    if st.button("Mark as Paid & Close Workflow", type="primary"):
                        mark_po_as_paid(po_data[0], pay_info)
                        st.success("પેમેન્ટ નોંધાઈ ગયું છે અને ફાઈલ ક્લોઝ થઈ ગઈ છે!")
                        st.rerun()
                with col_pay3:
                    st.write("")
                    confirm_delete_po = st.checkbox("Confirm delete", key=f"confirm_delete_po_{po_data[0]}")
                    if st.button("🗑️ Delete Entry Forever"):
                        if not confirm_delete_po:
                            st.warning("Please tick Confirm delete before deleting this payment entry.")
                        else:
                            delete_po(po_data[0])
                            st.error("ઓર્ડર કાયમ માટે રદ કરવામાં આવ્યો છે!")
                            st.rerun()

    # --- NEW SECTION: Upload Physical Signed Documents ---
    st.markdown("---")
    st.markdown("#### 📤 દસ્તાવેજ અપલોડ કરો (Upload Signed Documents to Vault)")
    st.info("સહી કરેલ PO, બિલ પેમેન્ટ અને પેસ્ટિંગ ફોર્મ અહીં કાયમી સાચવવા માટે અપલોડ કરો.")
    
    # 1. Get all Nondhs and map their subjects
    all_nondhs_for_vault = get_archives("All", "All")
    nondh_subject_map = {}
    
    # 2. Find which Nondhs are "Closed" (they already have a Signed Bill Pasting)
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("SELECT nondh_id FROM digital_vault WHERE doc_type = 'Signed Bill Pasting' AND nondh_id IS NOT NULL")
    closed_nondh_ids = [row[0] for row in c.fetchall()]
    conn.close()
    
    # 3. Create the filtered dropdown options (excluding closed Nondhs)
    nondh_opts = ["-- Select Nondh --"]
    for n in all_nondhs_for_vault:
        if len(n) == 4:
            n_id, n_subj = n[0], n[2]
            nondh_subject_map[n_id] = n_subj  # Save for the expander titles later
            if n_id not in closed_nondh_ids:
                nondh_opts.append(f"[{n_id}] {n_subj}") # Display ID and Subject
    
    col_u1, col_u2, col_u3 = st.columns(3)
    with col_u1:
        sel_nondh_vault = st.selectbox("કઈ નોંધ સાથે જોડવું છે? (Select Nondh ID)", nondh_opts)
    with col_u2:
        # Restricted Document Types
        sel_doc_type = st.selectbox("દસ્તાવેજનો પ્રકાર (Document Type)", [
            "Signed PO", "Signed Bill Payment", "Signed Bill Pasting"
        ])
    with col_u3:
        up_file = st.file_uploader("ફાઈલ પસંદ કરો (PDF/Image)", type=['pdf', 'jpg', 'jpeg', 'png'])
        
    if st.button("💾 વોલ્ટમાં સેવ કરો (Save to Permanent Vault)"):
        if sel_nondh_vault == "-- Select Nondh --":
            st.error("કૃપા કરીને નોંધ પસંદ કરો!")
        elif not up_file:
            st.error("કૃપા કરીને ફાઈલ અપલોડ કરો!")
        else:
            import re
            match = re.search(r'\[(\d+)\]', sel_nondh_vault)
            if match:
                n_id_val = int(match.group(1))
                file_bytes = up_file.getbuffer()
                save_file_to_vault(file_bytes, up_file.name, sel_doc_type, nondh_id=n_id_val, description="Manually Uploaded to Vault")
                st.success(f"{sel_doc_type} વોલ્ટમાં Nondh #{n_id_val} હેઠળ કાયમ માટે સેવ થઈ ગયું!")
                st.rerun()

    # --- UPDATED SECTION: Grouped Vault Display with Expanders ---
    st.markdown("---")
    st.markdown("#### 🔍 નોંધ મુજબ વોલ્ટ (Vault Grouped by Nondh ID)")
    
    current_year = datetime.date.today().year
    fy_options = ["All"] + [f"{y}-{str(y+1)[2:]}" for y in range(current_year-2, current_year+3)][::-1]
    
    col_f1, col_f2, col_f3 = st.columns(3)
    with col_f1: filter_fy = st.selectbox("નાણાકીય વર્ષ (Financial Year)", fy_options)
    with col_f2: filter_type = st.selectbox("પ્રકાર (Type)", ["All", "Signed Nondh", "Approval Letter", "Signed PO", "Party Invoice", "Signed Bill Payment", "Signed Bill Pasting", "Sadar Nondh Draft", "PO Draft", "Other"])
    with col_f3: search_kw = st.text_input("શબ્દથી શોધો (Search by Name/Tag)")

    vault_records = get_vault_files(filter_fy, filter_type, search_kw)
    
    if not vault_records: 
        st.info("કોઈ ડોક્યુમેન્ટ મળ્યા નથી.")
    else:
        st.success(f"કુલ {len(vault_records)} ડોક્યુમેન્ટ્સ મળ્યા.")
        
        # 1. Group records by Nondh ID
        grouped_vault = {}
        for record in vault_records:
            v_id, n_id, f_name, f_path, u_date, fy, month, d_type, desc = record
            group_key = n_id if n_id else "Unlinked"
            if group_key not in grouped_vault:
                grouped_vault[group_key] = []
            grouped_vault[group_key].append(record)
            
        # 2. Display them neatly under clickable Expanders
        for n_id_key, records in grouped_vault.items():
            
            # Determine the title of the expander
            if n_id_key == "Unlinked":
                expander_title = "📁 Unlinked / General Documents"
            else:
                subj = nondh_subject_map.get(n_id_key, "Unknown Subject")
                expander_title = f"📁 Nondh ID: {n_id_key} - {subj}"
            
            # Create a clickable dropdown panel (expander)
            # Create a clickable dropdown panel (expander)
            with st.expander(expander_title, expanded=False):
                for record in records:
                    v_id, n_id, f_name, f_path, u_date, fy, month, d_type, desc = record
                    with st.container(border=True):
                        col_info, col_btn1, col_btn2 = st.columns([6, 2, 2])
                        
                        with col_info:
                            st.markdown(f"**{d_type}**: {f_name}")
                            st.caption(f"🗓️ {u_date} | 📁 {fy} ({month})")
                            
                        with col_btn1:
                            # Use our intelligent cloud loader instead of direct os.path.exists
                            file_data = load_vault_file_bytes(f_path)
                            if file_data:
                                st.download_button("⬇️ Download", data=file_data, file_name=f_name, key=f"dl_vault_main_{v_id}")
                            else:
                                st.error("File missing completely")
                                
                        with col_btn2:
                            confirm_delete_vault = st.checkbox("Confirm", key=f"confirm_delete_vault_{v_id}")
                            if st.button("🗑️ Delete", key=f"del_vault_{v_id}"):
                                if not confirm_delete_vault:
                                    st.warning("Please tick Confirm before deleting this vault file.")
                                else:
                                    conn = sqlite3.connect(DB_FILE)
                                    c = conn.cursor()
                                    c.execute("DELETE FROM digital_vault WHERE id = ?", (v_id,))
                                    conn.commit()
                                    conn.close()
                                    push_db_to_github()
                                    if os.path.exists(f_path):
                                        try: os.remove(f_path)
                                        except: pass
                                    st.error(f"ફાઈલ '{f_name}' રદ કરવામાં આવી છે!")
                                    st.rerun()
